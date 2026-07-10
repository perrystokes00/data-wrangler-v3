"""
modules/dv_office_loader.py
============================
DataView v3 — Office Document Loader (Excel / Word).

Loaders parse .xlsx / .docx documents and capture the extracted rows into the
file_catalog.cat_* mirror tables via catalog_capture.capture() — they no longer
write dataview.dv_* directly. A dv_well header is NOT required at load time;
rows are keyed by the document UWI and promoted into dv_* later by
promote_catalog once a header exists.

Functions / capture targets:
    load_formation_tops_xlsx → cat_well_formation_top
    load_completion_xlsx     → cat_well_completion
    load_production_xlsx     → cat_prod_entity + cat_prod_volume (tall, per fluid)
    load_completion_docx     → cat_well_completion
    load_formation_tops_docx → cat_well_formation_top

Perforation and stimulation are outside the 11-table mirror scope (no
cat_well_perforation / cat_well_stimulation); their row counts are reported in
the result's "out_of_scope" note rather than captured.

All functions:
    engine : SQLAlchemy engine connected to DataView database
    path   : str — full file path on disk
    source : str — source code (default "DATA_LOADER")
    returns: {"loaded": int, "errors": [...], "wells": [...]}

Requirements:
    pip install openpyxl python-docx
"""
from __future__ import annotations

import hashlib
import re
import uuid
from datetime import datetime, date, timedelta
from pathlib import Path

import pandas as pd
from sqlalchemy import text


# =============================================================================
# HELPERS
# =============================================================================

def _uid() -> str:
    return hashlib.sha1(str(uuid.uuid4()).encode()).hexdigest()[:40]


def _trunc(v, n: int = 40) -> str | None:
    if v is None:
        return None
    s = str(v).strip()
    return s[:n] if s else None


def _safe_float(v) -> float | None:
    if v is None or str(v).strip() in ("", "None", "nan", "NaN"):
        return None
    try:
        return float(str(v).replace(",", "").strip())
    except (ValueError, TypeError):
        return None


def _excel_date(v) -> str | None:
    """Convert Excel serial date (float/int) or string date to ISO string."""
    if v is None:
        return None
    # Already a datetime or date
    if isinstance(v, (datetime, date)):
        return v.strftime("%Y-%m-%d")
    # Try float (Excel serial)
    try:
        serial = float(v)
        if 10000 < serial < 100000:  # plausible Excel date range
            dt = datetime(1899, 12, 30) + timedelta(days=serial)
            return dt.strftime("%Y-%m-%d")
    except (ValueError, TypeError):
        pass
    # Try string parse
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d-%b-%Y", "%B %d, %Y",
                "%b %d, %Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(str(v).strip(), fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return _trunc(v, 20)


def _result(loaded: int = 0, errors: list | None = None, **extra) -> dict:
    return {"loaded": loaded, "errors": errors or [], **extra}


def _well_exists(con, uwi: str) -> bool:
    """Legacy/unused — the catalog no longer checks dv_well existence."""
    return bool(str(uwi).strip()) if uwi else False


def _normalise_uwi(raw: str) -> str:
    """Strip dashes from UWI — 42-317-12345-00-00 → 42317123450000 — AND repair
    Excel float/scientific-notation corruption. A 14-digit UWI read from a numeric
    xlsx/csv cell can arrive as '4.23171E+13' or '42317123450000.0'; both must be
    restored to the exact digit string or the promote join to dv_well fails."""
    return _uwi_fix(raw)


def _uwi_fix(raw) -> str:
    """Repair a UWI that Excel/pandas may have mangled into float or scientific
    notation, then strip separators. Returns '' for blanks."""
    if raw is None:
        return ""
    s = str(raw).strip()
    if not s or s.lower() in ("none", "nan"):
        return ""
    # scientific notation (4.23171E+13) or trailing-.0 float (42317123450000.0)
    import re as _re
    if _re.fullmatch(r"[+-]?\d+(\.\d+)?[eE][+-]?\d+", s) or _re.fullmatch(r"\d+\.0+", s):
        try:
            s = format(int(float(s)), "d")   # 4.23171E+13 -> 42317123450000
        except Exception:
            pass
    # strip dashes/spaces used in display UWIs
    return _re.sub(r"[-\s]", "", s)


def _find_well(con, raw_uwi: str) -> str | None:
    """
    Return the document's own UWI (trimmed).

    The catalog no longer requires the well to pre-exist in dv_well — documents
    are loaded before the well header exists, so rows are keyed by the document
    UWI and reconciled to the master later. `con` is accepted for signature
    stability but is not queried.
    """
    key = str(raw_uwi).strip() if raw_uwi else ""
    return key or None


# Resilient import: works whether catalog_capture lands in modules/ or root.
try:
    from modules.catalog_capture import capture
except ImportError:
    from catalog_capture import capture


def _ts() -> str:
    """Plain-value timestamp for capture (cat_* rows hold literal values)."""
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _inv_for(engine, path: str):
    """Resolve INVENTORY_ID for provenance from the file catalog, or None."""
    try:
        with engine.connect() as c:
            r = c.execute(text(
                "SELECT TOP 1 INVENTORY_ID FROM file_catalog.GLOBAL_FILE_CATALOG "
                "WHERE FILE_PATH = :p"), {"p": path}).fetchone()
        return r[0] if r else None
    except Exception:
        return None


# =============================================================================
# 1. FORMATION TOPS FROM XLSX → dv_well_formation_top
# =============================================================================

def _read_tabular(path, sheet_keywords=None):
    """Read a tabular document as a string-typed DataFrame. .csv -> read_csv;
    .xlsx/.xls -> read_excel (optionally picking a sheet whose name matches one of
    sheet_keywords, else the first). dtype=str so UWIs never become floats."""
    ext = Path(path).suffix.lower()
    if ext == ".csv":
        return pd.read_csv(path, dtype=str, keep_default_na=False)
    xl = pd.ExcelFile(path)
    sheet = xl.sheet_names[0]
    if sheet_keywords:
        sheet = next((s for s in xl.sheet_names
                      if any(k in s.lower() for k in sheet_keywords)),
                     xl.sheet_names[0])
    return pd.read_excel(path, sheet_name=sheet, dtype=str)


def load_formation_tops_xlsx(engine, path: str,
                              source: str = "DATA_LOADER") -> dict:
    """
    Load formation tops from Formation_Tops.xlsx style file.

    Expected columns (case-insensitive):
        UWI, WELL_NAME, FORMATION, TOP_MD_FT, BASE_MD_FT,
        TOP_TVD_FT (optional), NET_PAY_FT (optional),
        FLUID (optional), PICKED_BY (optional), PICK_DATE (optional)
    """
    errors  = []
    loaded  = 0
    wells   = []

    try:
        xl = pd.ExcelFile(path)
        # Find the formation tops sheet
        sheet = next(
            (s for s in xl.sheet_names
             if any(kw in s.lower() for kw in
                    ("formation", "top", "pick", "strat"))),
            xl.sheet_names[0]
        )
        df = pd.read_excel(path, sheet_name=sheet, dtype=str)
        df.columns = [c.upper().strip() for c in df.columns]
        df = df.dropna(how="all")
    except Exception as exc:
        return _result(errors=[f"Could not read Excel: {exc}"])

    # Column name aliases
    _COL = {
        "UWI":         next((c for c in df.columns if "UWI" in c), None),
        "WELL_NAME":   next((c for c in df.columns if "WELL" in c and "NAME" in c), None),
        "FORMATION":   next((c for c in df.columns if any(k in c for k in ("FORMATION","FORM","UNIT","STRAT"))), None),
        "TOP_MD":      next((c for c in df.columns if "TOP" in c and ("MD" in c or "FT" in c) and "BASE" not in c and "TVD" not in c), None),
        "BASE_MD":     next((c for c in df.columns if "BASE" in c and ("MD" in c or "FT" in c)), None),
        "TOP_TVD":     next((c for c in df.columns if "TVD" in c and "BASE" not in c), None),
        "NET_PAY":     next((c for c in df.columns if "NET" in c and "PAY" in c), None),
        "FLUID":       next((c for c in df.columns if "FLUID" in c), None),
        "PICKED_BY":   next((c for c in df.columns if "PICK" in c and "BY" in c), None),
        "PICK_DATE":   next((c for c in df.columns if "PICK" in c and "DATE" in c), None),
    }

    if not _COL["UWI"] or not _COL["FORMATION"] or not _COL["TOP_MD"]:
        return _result(errors=[
            f"Required columns missing. Found: {list(df.columns)}. "
            f"Need: UWI, FORMATION, TOP_MD_FT"
        ])

    ts  = _ts()
    inv = _inv_for(engine, path)
    out = []
    for i, row in df.iterrows():
        raw_uwi    = _uwi_fix(row.get(_COL["UWI"], ""))
        form_name  = _trunc(row.get(_COL["FORMATION"]), 255)
        top_depth  = _safe_float(row.get(_COL["TOP_MD"]))
        base_depth = _safe_float(row.get(_COL["BASE_MD"])) if _COL["BASE_MD"] else None

        if not raw_uwi or not form_name or top_depth is None:
            continue

        out.append({
            "uwi":              raw_uwi,
            "strat_unit_id":    _uid(),
            "interp_id":        _uid(),
            "strat_unit_name":  form_name,
            "strat_unit_type":  "FORMATION",
            "top_depth":        top_depth,
            "base_depth":       base_depth,
            "depth_ouom":       "ft",
            "active_ind":       "Y",
            "row_created_by":   "DataWrangler",
            "row_created_date": ts,
        })
        if raw_uwi not in wells:
            wells.append(raw_uwi)

    try:
        loaded = capture(engine, "cat_well_formation_top", out,
                         uwi=None, inventory_id=inv,
                         source_path=path, source=source)
    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=wells)


# =============================================================================
# 2. COMPLETION PARAMETERS FROM XLSX
#    → dv_well_completion + dv_well_perforation + dv_well_stimulation
# =============================================================================

def load_completion_xlsx(engine, path: str,
                         source: str = "DATA_LOADER") -> dict:
    """
    Load completion data from Completion_Parameters.xlsx style file.

    Reads three sheets:
        Completions  → dv_well_completion + dv_well_stimulation
        Perforations → dv_well_perforation
        Well Header  → UPDATE dv_well (lat/lon/operator if blank)
    """
    errors = []
    loaded = 0
    wells  = []

    try:
        xl = pd.ExcelFile(path)
        sheets = {s.lower(): s for s in xl.sheet_names}
    except Exception as exc:
        return _result(errors=[f"Could not read Excel: {exc}"])

    def _read_sheet(keywords):
        sheet = next(
            (sheets[k] for k in sheets
             if any(kw in k for kw in keywords)),
            None
        )
        if not sheet:
            return pd.DataFrame()
        df = pd.read_excel(path, sheet_name=sheet, dtype=str)
        df.columns = [c.upper().strip() for c in df.columns]
        return df.dropna(how="all")

    comp_df = _read_sheet(["complet", "frac", "stim"])
    perf_df = _read_sheet(["perf"])
    hdr_df  = _read_sheet(["header", "well header", "wells"])

    ts  = _ts()
    inv = _inv_for(engine, path)
    out = []
    for i, row in comp_df.iterrows():
        raw_uwi = _uwi_fix(row.get("UWI", ""))
        if not raw_uwi:
            continue
        out.append({
            "uwi":              raw_uwi,
            "completion_id":    _uid(),
            "completion_type":  "HYDRAULIC_FRACTURE",
            "completion_date":  _excel_date(row.get("COMP_DATE")),
            "top_depth":        _safe_float(row.get("PERF_TOP_FT")),
            "base_depth":       _safe_float(row.get("PERF_BOT_FT")),
            "depth_ouom":       "ft",
            "active_ind":       "Y",
            "row_created_by":   "DataWrangler",
            "row_created_date": ts,
        })
        if raw_uwi not in wells:
            wells.append(raw_uwi)

    try:
        loaded = capture(engine, "cat_well_completion", out,
                         uwi=None, inventory_id=inv,
                         source_path=path, source=source)
    except Exception as exc:
        return _result(errors=[str(exc)])

    # Perforation / stimulation are outside the 11-table mirror scope
    # (no cat_well_perforation / cat_well_stimulation). Report the count so
    # the data isn't silently dropped.
    n_perf = int(len(perf_df))
    n_stim = int(len(comp_df))
    note = None
    if n_perf or n_stim:
        note = (f"perforation ({n_perf}) and stimulation ({n_stim}) rows not "
                f"captured — outside catalog mirror scope")

    return _result(loaded=loaded, errors=errors, wells=wells,
                   out_of_scope=note)


# =============================================================================
# 3. PRODUCTION DATA FROM XLSX → dv_prod_entity + dv_prod_volume
# =============================================================================

def load_production_xlsx(engine, path: str,
                         source: str = "DATA_LOADER") -> dict:
    """
    Load monthly production from Production_Data.xlsx style file.

    Expected columns:
        UWI, WELL_NAME, DATE, OIL_BBL, GAS_MCF, WATER_BBL,
        BOE (optional), TUBING_PRESS_PSI (optional), STATUS (optional)
    """
    errors = []
    loaded = 0
    wells  = []

    try:
        df = _read_tabular(path, sheet_keywords=("prod", "monthly", "volume"))
        df.columns = [c.upper().strip() for c in df.columns]
        df = df.dropna(how="all")
    except Exception as exc:
        return _result(errors=[f"Could not read Excel: {exc}"])

    ts  = _ts()
    inv = _inv_for(engine, path)
    entities: dict[str, str] = {}   # uwi -> prod_entity_id
    ent_rows = []
    vol_rows = []

    for i, row in df.iterrows():
        raw_uwi = _uwi_fix(row.get("UWI", ""))
        if not raw_uwi:
            continue

        if raw_uwi not in entities:
            eid = _uid()
            entities[raw_uwi] = eid
            ent_rows.append({
                "uwi":              raw_uwi,
                "prod_entity_id":   eid,
                "prod_entity_type": "WELL",
                "prod_entity_name": _trunc(row.get("WELL_NAME") or raw_uwi, 255),
                "primary_fluid":    "OIL",
                "active_ind":       "Y",
                "row_created_by":   "DataWrangler",
                "row_created_date": ts,
            })
            if raw_uwi not in wells:
                wells.append(raw_uwi)
        eid = entities[raw_uwi]

        prod_date = _excel_date(row.get("DATE"))
        if not prod_date:
            errors.append(f"Row {i+1}: could not parse date '{row.get('DATE')}'")
            continue
        try:
            dt = datetime.strptime(prod_date, "%Y-%m-%d")
            period = f"{dt.year:04d}-{dt.month:02d}"
        except Exception:
            period = (prod_date or "")[:7]
        # period_date is NOT NULL in dv_prod_volume — never write a volume row
        # without a valid YYYY-MM period (a blank/None here is what broke promote).
        if not period or len(period) < 7 or not period[:4].isdigit():
            errors.append(f"Row {i+1}: no usable period from date '{prod_date}'")
            continue

        for fluid, key, uom in (("OIL", "OIL_BBL", "BBL"),
                                ("GAS", "GAS_MCF", "MCF"),
                                ("WATER", "WATER_BBL", "BBL")):
            vol = _safe_float(row.get(key))
            if vol is None:
                continue
            vol_rows.append({
                "uwi":              raw_uwi,     # cat_prod_volume UWI helper
                "prod_entity_id":   eid,
                "period_date":      period,
                "fluid_type":       fluid,
                "volume":           vol,
                "volume_ouom":      uom,
                "active_ind":       "Y",
                "row_created_by":   "DataWrangler",
                "row_created_date": ts,
            })

    # Collapse duplicate production records before capture. dv_prod_volume's PK
    # is (prod_entity_id, period_date, fluid_type); a source sheet can list the
    # same well/month/fluid more than once (re-exports, merged files), which
    # would later collide on promote. Keep the LAST occurrence per key — last
    # write wins, matching how a corrected re-export should supersede an earlier
    # row. This makes the capture idempotent at the key level.
    if vol_rows:
        _seen = {}
        for vr in vol_rows:
            _seen[(vr["prod_entity_id"], vr["period_date"],
                   vr["fluid_type"])] = vr
        _dropped = len(vol_rows) - len(_seen)
        if _dropped:
            errors.append(f"Collapsed {_dropped} duplicate production "
                          f"record(s) (same entity/period/fluid).")
        vol_rows = list(_seen.values())

    try:
        capture(engine, "cat_prod_entity", ent_rows,
                uwi=None, inventory_id=inv, source_path=path, source=source)
        loaded = capture(engine, "cat_prod_volume", vol_rows,
                         uwi=None, inventory_id=inv,
                         source_path=path, source=source)
    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=wells)


# =============================================================================
# 4. WELL COMPLETION REPORT FROM DOCX
#    → dv_well_completion + dv_well_perforation + dv_well_stimulation
# =============================================================================

def load_completion_docx(engine, path: str,
                         source: str = "DATA_LOADER") -> dict:
    """
    Load completion report from a Word .docx file.
    Extracts tables by heading context:
        Well Identification → well header fields
        Perforation Intervals → dv_well_perforation
        Stimulation Parameters → dv_well_stimulation
        Initial Production → remark on dv_well_completion
    """
    errors = []
    loaded = 0

    try:
        from docx import Document
    except ImportError:
        return _result(errors=["python-docx not installed: pip install python-docx"])

    try:
        doc = Document(path)
    except Exception as exc:
        return _result(errors=[f"Could not read Word doc: {exc}"])

    # ── Extract all tables with their preceding heading ────────────────────
    def _table_to_df(table):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if not rows:
            return pd.DataFrame()
        # Detect if first row is a header
        first = rows[0]
        if len(first) == 2 and all(first):
            # Key-value table
            return pd.DataFrame(rows, columns=["FIELD", "VALUE"])
        if len(rows) > 1:
            return pd.DataFrame(rows[1:], columns=[c.upper() for c in rows[0]])
        return pd.DataFrame()

    # Walk document body to get heading → table pairs
    sections: dict[str, list] = {}
    current_heading = "GENERAL"
    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            style = block.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
            )
            if style is not None:
                sval = style.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
                )
                if "heading" in sval.lower() or "Heading" in sval:
                    texts = [n.text or "" for n in block.iter()
                             if n.text and n.tag.endswith("t")]
                    current_heading = " ".join(texts).strip().upper()
        elif tag == "tbl":
            from docx.oxml.ns import qn
            from docx.table import Table as DocxTable
            tbl_obj = DocxTable(block, doc)
            df = _table_to_df(tbl_obj)
            if current_heading not in sections:
                sections[current_heading] = []
            sections[current_heading].append(df)

    # ── Extract UWI from Well Identification table ─────────────────────────
    uwi       = None
    comp_date = None
    well_info = {}

    for heading, tables in sections.items():
        if any(k in heading for k in ("WELL ID", "IDENTIFICATION", "WELL IDENT")):
            for df in tables:
                if "FIELD" in df.columns and "VALUE" in df.columns:
                    for _, r in df.iterrows():
                        field = str(r["FIELD"]).upper().strip()
                        val   = str(r["VALUE"]).strip()
                        if not val or val.lower() in ("none", ""):
                            continue
                        if "UWI" in field:
                            uwi = val
                        elif "COMPLETION DATE" in field:
                            comp_date = _excel_date(val)
                        elif "SPUD" in field:
                            well_info["SPUD_DATE"] = _excel_date(val)
                        elif "TOTAL DEPTH" in field:
                            nums = re.findall(r"[\d,]+", val.split("/")[0])
                            if nums:
                                well_info["FINAL_TD"] = float(nums[0].replace(",",""))
                        elif "KB ELEVATION" in field:
                            nums = re.findall(r"[\d,]+", val)
                            if nums:
                                well_info["KB_ELEVATION"] = float(nums[0].replace(",",""))

    if not uwi:
        return _result(errors=["Could not find UWI in document"])

    ts  = _ts()
    inv = _inv_for(engine, path)
    db_uwi = _find_well(None, uwi)
    if not db_uwi:
        return _result(errors=[f"UWI '{uwi}' is empty"])

    comp_row = [{
        "uwi":              db_uwi,
        "completion_id":    _uid(),
        "completion_type":  "HYDRAULIC_FRACTURE",
        "completion_date":  comp_date,
        "active_ind":       "Y",
        "row_created_by":   "DataWrangler",
        "row_created_date": ts,
    }]
    try:
        loaded = capture(engine, "cat_well_completion", comp_row,
                         uwi=None, inventory_id=inv,
                         source_path=path, source=source)
    except Exception as exc:
        return _result(errors=[str(exc)])

    # Perforation / stimulation are outside the mirror scope — count them so
    # the data isn't silently dropped.
    n_perf = sum(
        len(df) for h, tabs in sections.items()
        if any(k in h for k in ("PERF", "INTERVAL")) for df in tabs)
    n_stim = sum(
        1 for h in sections if any(k in h for k in ("STIM", "FRAC", "TREAT")))
    note = None
    if n_perf or n_stim:
        note = (f"perforation ({n_perf}) and stimulation ({n_stim}) not "
                f"captured — outside catalog mirror scope")

    return _result(loaded=loaded, errors=errors,
                   wells=[db_uwi], out_of_scope=note)


# =============================================================================
# 5. FORMATION TOPS FROM DOCX → dv_well_formation_top
# =============================================================================

def load_formation_tops_docx(engine, path: str,
                              source: str = "DATA_LOADER") -> dict:
    """
    Load formation tops from a Word .docx file.
    Finds tables with UWI/Formation/Top MD columns.
    """
    errors = []
    loaded = 0
    wells  = []

    try:
        from docx import Document
    except ImportError:
        return _result(errors=["python-docx not installed: pip install python-docx"])

    try:
        doc = Document(path)
    except Exception as exc:
        return _result(errors=[f"Could not read Word doc: {exc}"])

    ts  = _ts()
    inv = _inv_for(engine, path)
    out = []
    try:
        for ti, table in enumerate(doc.tables):
            if not table.rows:
                continue
            header = [cell.text.strip().upper()
                      for cell in table.rows[0].cells]
            cols = {h: i for i, h in enumerate(header)}

            uwi_col  = next((cols[c] for c in cols if "UWI" in c), None)
            form_col = next((cols[c] for c in cols
                             if any(k in c for k in
                                    ("FORMATION", "FORM", "UNIT", "STRAT"))), None)
            top_col  = next((cols[c] for c in cols
                             if "TOP" in c and "BASE" not in c
                             and "TVD" not in c), None)
            base_col = next((cols[c] for c in cols if "BASE" in c), None)

            if uwi_col is None or form_col is None or top_col is None:
                continue  # not a formation tops table

            for row in table.rows[1:]:
                cells   = [cell.text.strip() for cell in row.cells]
                raw_uwi = _uwi_fix(cells[uwi_col] if uwi_col < len(cells) else "")
                if not raw_uwi:
                    continue
                form_name  = _trunc(cells[form_col], 255) if form_col < len(cells) else None
                top_depth  = _safe_float(cells[top_col]) if top_col < len(cells) else None
                base_depth = _safe_float(cells[base_col]) if base_col and base_col < len(cells) else None
                if not form_name or top_depth is None:
                    continue
                out.append({
                    "uwi":              raw_uwi,
                    "strat_unit_id":    _uid(),
                    "interp_id":        _uid(),
                    "strat_unit_name":  form_name,
                    "strat_unit_type":  "FORMATION",
                    "top_depth":        top_depth,
                    "base_depth":       base_depth,
                    "depth_ouom":       "ft",
                    "active_ind":       "Y",
                    "row_created_by":   "DataWrangler",
                    "row_created_date": ts,
                })
                if raw_uwi not in wells:
                    wells.append(raw_uwi)

        loaded = capture(engine, "cat_well_formation_top", out,
                         uwi=None, inventory_id=inv,
                         source_path=path, source=source)
    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=wells)


# =============================================================================
# DISPATCHER — call the right loader based on file extension + doc type hint
# =============================================================================

DOC_TYPE_LOADERS = {
    # (extension, doc_type_hint) → function
    ("xlsx", "formation"):   load_formation_tops_xlsx,
    ("xlsx", "completion"):  load_completion_xlsx,
    ("xlsx", "production"):  load_production_xlsx,
    ("docx", "completion"):  load_completion_docx,
    ("docx", "formation"):   load_formation_tops_docx,
}


def dispatch(engine, path: str, doc_type_hint: str = "",
             source: str = "DATA_LOADER") -> dict:
    """
    Automatically select and call the right loader based on
    file extension and doc_type_hint keyword.

    doc_type_hint examples: "formation", "completion", "production"
    """
    ext  = Path(path).suffix.lower().lstrip(".")
    hint = doc_type_hint.lower()

    # Match most specific key first
    key = next(
        ((e, h) for (e, h) in DOC_TYPE_LOADERS
         if e == ext and h in hint),
        None
    )
    if key:
        return DOC_TYPE_LOADERS[key](engine, path, source)

    # Fallback: guess from filename
    fname = Path(path).stem.lower()
    if "formation" in fname or "tops" in fname or "pick" in fname:
        if ext == "xlsx":
            return load_formation_tops_xlsx(engine, path, source)
        if ext == "docx":
            return load_formation_tops_docx(engine, path, source)
    if "complet" in fname or "frac" in fname or "stim" in fname:
        if ext == "xlsx":
            return load_completion_xlsx(engine, path, source)
        if ext == "docx":
            return load_completion_docx(engine, path, source)
    if "prod" in fname or "volume" in fname or "monthly" in fname:
        if ext in ("xlsx", "csv"):
            return load_production_xlsx(engine, path, source)
    # generic csv fallback: a csv with UWI-ish columns is production data
    if ext == "csv":
        return load_production_xlsx(engine, path, source)

    return _result(errors=[
        f"No loader found for extension='{ext}' hint='{doc_type_hint}'. "
        f"Supported: formation/xlsx, completion/xlsx, production/xlsx, "
        f"formation/docx, completion/docx"
    ])
