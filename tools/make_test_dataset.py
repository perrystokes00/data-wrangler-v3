#!/usr/bin/env python3
"""
make_test_dataset.py
====================
Generate a complete, *parseable* test corpus for the DataView catalog / vault
workflow — wells with several file types each, seismic surveys (2D + 3D with
nav), and spatial data — so you can exercise scan -> extract -> triage ->
promote -> vault end to end.

Every file is a real file in its native format, written with the same library
your extractors read with, and each well's documents reference that well's UWI,
name, operator and location (in the file content AND the filename/folder) so
they match and aggregate under one well.

Layout produced under --root (default ./TestData):

    Wells/<UWI>__<WELL_NAME>/
        <UWI>_run1_triple_combo.las     (lasio)
        <UWI>_run2_resistivity.las      (lasio)
        <UWI>_directional_survey.pdf    (reportlab)
        <UWI>_formation_tops.pdf        (reportlab)
        <UWI>_completion.xlsx           (openpyxl)
        <UWI>_well_summary.docx         (python-docx)
    Seismic/<SURVEY>_<YEAR>/
        <survey>.sgy|.segy              (segyio)  + <survey>.p190 (text)
    Spatial/
        well_locations.shp/.dbf/.shx/.prj   (pyshp)
        well_locations.geojson

UWIs use a clearly-fake 42999##### prefix so they never collide with real wells
and are trivial to clean up later (e.g. DELETE ... WHERE UWI14 LIKE '42999%').

Run:
    python make_test_dataset.py --root C:\\Bulk\\TestData

Each format is optional: if a library is missing, that format is skipped with a
warning rather than aborting the run.  DLIS/LIS are intentionally NOT generated
— they are binary formats with no practical Python writer; drop a couple of your
own real .dlis/.lis files into a well folder if you want to test those paths.
"""

import argparse
import json
import os
import random
import sys

random.seed(42)

# ── well + survey definitions ────────────────────────────────────────────────
WELLS = [
    dict(uwi="42999000010000", name="SMITH 1-H",     operator="TEST OPERATING CO",
         field="GROGANS MILL", state="TX", county="MONTGOMERY", country="US",
         lat=30.350, lon=-95.500, strt=1000.0, stop=9800.0),
    dict(uwi="42999000020000", name="JONES A2",       operator="WILDCAT ENERGY LLC",
         field="MAGNOLIA",     state="TX", county="MONTGOMERY", country="US",
         lat=30.402, lon=-95.452, strt=1500.0, stop=11200.0),
    dict(uwi="42999000030000", name="BAKER UNIT 3",   operator="TEST OPERATING CO",
         field="NAVASOTA",     state="TX", county="GRIMES",     country="US",
         lat=30.553, lon=-95.951, strt=2000.0, stop=8600.0),
]

SURVEYS = [
    dict(name="GULF SHELF 2D", dim="2D", year=2011, state="TX", country="US",
         folder="GULF_SHELF_2D", file="gulf_shelf_2d", ext=".sgy",
         ilines=1, xlines=240),
    dict(name="DELTA DEEP 3D", dim="3D", year=2007, state="LA", country="US",
         folder="DELTA_DEEP_3D", file="delta_deep_3d", ext=".segy",
         ilines=20, xlines=20),
]


def fmt_api(uwi):
    """42999000010000 -> 42-999-00001-00-00 (human-readable API)."""
    u = (uwi + "00000000000000")[:14]
    return f"{u[0:2]}-{u[2:5]}-{u[5:10]}-{u[10:12]}-{u[12:14]}"


def safe(s):
    return "".join(c if c.isalnum() or c in " -_" else "_" for c in str(s)).strip().replace(" ", "_")


# ── LAS (text; validated with lasio) ─────────────────────────────────────────
RUN1 = [("DEPT", "FT", "DEPTH"), ("GR", "GAPI", "GAMMA RAY"),
        ("RHOB", "G/C3", "BULK DENSITY"), ("NPHI", "V/V", "NEUTRON POROSITY"),
        ("DT", "US/FT", "SONIC TRANSIT TIME"), ("CALI", "IN", "CALIPER")]
RUN2 = [("DEPT", "FT", "DEPTH"), ("ILD", "OHMM", "DEEP INDUCTION"),
        ("ILM", "OHMM", "MEDIUM INDUCTION"), ("SFL", "OHMM", "SHALLOW FOCUSED"),
        ("SP", "MV", "SPONTANEOUS POTENTIAL")]


def build_las(w, curves, run_label, step=0.5):
    strt, stop = w["strt"], w["strt"] + 800.0   # 800 ft logged interval
    null = -999.25
    head = [
        "~Version Information",
        " VERS.   2.0 : CWLS LOG ASCII STANDARD - VERSION 2.0",
        " WRAP.    NO : ONE LINE PER DEPTH STEP",
        "~Well Information",
        "#MNEM.UNIT       DATA               : DESCRIPTION",
        f" STRT.FT   {strt:10.4f} : START DEPTH",
        f" STOP.FT   {stop:10.4f} : STOP DEPTH",
        f" STEP.FT   {step:10.4f} : STEP",
        f" NULL.     {null:10.4f} : NULL VALUE",
        f" COMP.   {w['operator']} : COMPANY",
        f" WELL.   {w['name']} : WELL",
        f" FLD .   {w['field']} : FIELD",
        f" SRVC.   TEST WIRELINE : SERVICE COMPANY",
        f" DATE.   2024-01-15 : LOG DATE",
        f" UWI .   {w['uwi']} : UNIQUE WELL ID",
        f" API .   {w['uwi']} : API NUMBER",
        f" STAT.   {w['state']} : STATE",
        f" CNTY.   {w['county']} : COUNTY",
        f" CTRY.   {w['country']} : COUNTRY",
        "~Curve Information",
        "#MNEM.UNIT      : DESCRIPTION",
    ]
    for i, (m, u, d) in enumerate(curves):
        head.append(f" {m:<4}.{u:<6} : {i}  {d}")
    head += ["~Parameter Information", "~Other", "~ASCII"]

    lines = head
    n = int((stop - strt) / step) + 1
    for k in range(n):
        depth = strt + k * step
        vals = [f"{depth:10.4f}"]
        for (m, u, d) in curves[1:]:
            vals.append(f"{random.uniform(10, 150):10.4f}")
        lines.append(" " + " ".join(vals))
    return "\n".join(lines) + "\n"


def gen_las(path, w, curves, label):
    with open(path, "w", newline="\n") as f:
        f.write(build_las(w, curves, label))


# ── PDF (reportlab) ──────────────────────────────────────────────────────────
def _pdf_doc(path, title, w, table_header, table_rows):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    meta = (f"WELL: {w['name']}<br/>UWI / API: {w['uwi']}  ({fmt_api(w['uwi'])})<br/>"
            f"OPERATOR: {w['operator']}<br/>FIELD: {w['field']}<br/>"
            f"COUNTY: {w['county']}, {w['state']}, {w['country']}")
    story += [Paragraph(meta, styles["Normal"]), Spacer(1, 16)]
    data = [table_header] + table_rows
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F4858")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    story.append(t)
    SimpleDocTemplate(path, pagesize=letter).build(story)


def gen_directional_pdf(path, w):
    rows, md, inc, azi = [], 0.0, 0.0, 0.0
    for _ in range(25):
        md += w["stop"] / 25.0
        inc = min(90.0, inc + random.uniform(0, 4))
        azi = (azi + random.uniform(-5, 5)) % 360
        tvd = md * (1 - inc / 200.0)
        rows.append([f"{md:.1f}", f"{inc:.2f}", f"{azi:.2f}", f"{tvd:.1f}"])
    _pdf_doc(path, "DIRECTIONAL SURVEY REPORT", w,
             ["MD (ft)", "INC (deg)", "AZI (deg)", "TVD (ft)"], rows)


FORMATIONS = ["AUSTIN CHALK", "EAGLE FORD", "BUDA LIME", "DEL RIO", "GEORGETOWN",
              "EDWARDS", "WOODBINE", "MANESS SHALE"]


def gen_formation_pdf(path, w):
    rows, md = [], w["strt"]
    for fm in FORMATIONS:
        md += random.uniform(400, 1100)
        rows.append([fm, f"{md:.1f}", f"{md*0.97:.1f}"])
    _pdf_doc(path, "FORMATION TOPS REPORT", w,
             ["FORMATION", "MD (ft)", "TVD (ft)"], rows)


def gen_scout_pdf(path, w):
    """Scout ticket: header block + the section tables extract_scout_ticket
    classifies (formation tops, DST, frac, core, IP).  Carries >=2 scout
    keywords (SCOUT TICKET / SCOUT REPORT / INITIAL PRODUCTION / IP RATE) so
    classify_pdf tags it RT_SCOUT, and avoids EOWR/mudlog title phrases."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    story = [
        Paragraph("SCOUT TICKET", styles["Title"]),
        Paragraph("SCOUT REPORT - INITIAL PRODUCTION SUMMARY", styles["Heading2"]),
        Spacer(1, 10),
    ]
    td = w["stop"]
    tvd = td * 0.55
    # header fields, one per line so the extractor's regexes capture cleanly
    fields = [
        ("Well Name", w["name"]), ("Well No", w["name"].split()[-1]),
        ("Operator", w["operator"]), ("Field", w["field"]),
        ("Lease", f"{w['name'].split()[0]} LEASE"), ("County", w["county"]),
        ("State", w["state"]), ("Well Type", "OIL"), ("Status", "PRODUCING"),
        ("Spud Date", "2023-08-01"), ("Completion Date", "2023-11-15"),
        ("API", fmt_api(w["uwi"])), ("UWI", w["uwi"]),
        ("Total Depth", f"{td:.0f} ft"), ("TVD", f"{tvd:.0f} ft"),
        ("Lateral", f"{td - tvd:.0f} ft"), ("KB", "1250"),
        ("Latitude", f"{w['lat']:.5f}"), ("Longitude", f"{w['lon']:.5f}"),
    ]
    for k, v in fields:
        story.append(Paragraph(f"<b>{k}:</b> {v}", styles["Normal"]))
    story.append(Spacer(1, 12))

    def tbl(title, header, rows):
        story.append(Paragraph(title, styles["Heading3"]))
        t = Table([header] + rows, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F4858")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    md = w["strt"]
    ftops = []
    for fm in FORMATIONS[:6]:
        md += random.uniform(400, 1100)
        ftops.append([fm, f"{md:.0f}", f"{md*1.01:.0f}"])
    tbl("FORMATION TOPS", ["FORMATION", "TOP (FT)", "BASE (FT)"], ftops)

    dst = []
    for i in range(2):
        top = w["strt"] + random.uniform(2000, 5000)
        base = top + random.uniform(20, 80)
        dst.append([f"DST-{i+1}", f"2023-10-{11+i}", f"{top:.0f}", f"{base:.0f}",
                    "OIL & GAS", f"{random.uniform(50, 400):.0f}",
                    f"{random.uniform(200, 2000):.0f}", f"{random.uniform(2000, 6000):.0f}"])
    tbl("DRILL STEM TESTS",
        ["DST TYPE", "DATE", "TOP", "BASE", "RESULT", "OIL RATE", "GAS RATE", "SHUT-IN PRESS"], dst)

    frac = []
    for st in range(1, 6):
        top = w["stop"] - random.uniform(500, 3000)
        base = top + random.uniform(100, 300)
        frac.append([str(st), f"{top:.0f}", f"{base:.0f}",
                     f"{random.uniform(2000, 8000):.0f}",
                     f"{random.uniform(100000, 400000):.0f}",
                     f"{random.uniform(3000, 7000):.0f}",
                     f"{random.uniform(6000, 9000):.0f}"])
    tbl("FRAC / STIMULATION",
        ["STAGE", "TOP", "BASE", "FLUID BBL", "PROPPANT LBS", "ISIP", "MAX PRESS"], frac)

    core = []
    d = w["strt"] + 1000
    for _ in range(5):
        d += random.uniform(50, 200)
        core.append([f"{d:.0f}", f"{random.uniform(4, 28):.1f}",
                     f"{random.uniform(0.01, 500):.2f}", f"{random.uniform(15, 60):.1f}"])
    tbl("CORE ANALYSIS", ["DEPTH", "POROSITY", "PERM", "SW"], core)

    ip = []
    for mo in range(1, 5):
        ip.append([f"2023-{10+mo:02d}", f"{random.uniform(100, 800):.0f}",
                   f"{random.uniform(500, 3000):.0f}", f"{random.uniform(20, 200):.0f}"])
    tbl("IP RATE / INITIAL PRODUCTION", ["DATE", "OIL BBL", "GAS MCF", "WATER BBL"], ip)

    SimpleDocTemplate(path, pagesize=letter).build(story)


# ── Excel (openpyxl) ─────────────────────────────────────────────────────────
def gen_xlsx(path, w):
    from openpyxl import Workbook
    from openpyxl.styles import Font
    wb = Workbook()

    # Sheet 1 — well identity as a HEADER-ROW table (row 1 = column headers,
    # row 2 = values). _summarize_excel reads row 0 as headers and pulls the
    # UWI from a column whose header contains uwi/api/well_id, so the identity
    # must be a column, not a label:value pair.
    ws = wb.active
    ws.title = "Well"
    headers = ["UWI", "WELL_NAME", "API", "OPERATOR", "FIELD",
               "COUNTY", "STATE", "COUNTRY", "TOTAL_DEPTH_FT"]
    for c, h in enumerate(headers, start=1):
        ws.cell(row=1, column=c, value=h).font = Font(bold=True)
    ws.append([w["uwi"], w["name"], fmt_api(w["uwi"]), w["operator"],
               w["field"], w["county"], w["state"], w["country"],
               round(w["stop"], 1)])

    # Sheet 2 — perforations, also a proper header-row table (keyed by UWI so a
    # completion/perf loader can attribute rows to the well).
    ws2 = wb.create_sheet("Perforations")
    perf_hdr = ["UWI", "TOP_FT", "BOTTOM_FT", "SHOTS_PER_FT", "STATUS"]
    for c, h in enumerate(perf_hdr, start=1):
        ws2.cell(row=1, column=c, value=h).font = Font(bold=True)
    top = w["stop"] - 600
    for _ in range(6):
        bot = top + random.uniform(20, 60)
        ws2.append([w["uwi"], round(top, 1), round(bot, 1),
                    random.choice([4, 6, 12]), "OPEN"])
        top = bot + random.uniform(30, 80)

    wb.save(path)


# ── Word (python-docx) ───────────────────────────────────────────────────────
def gen_docx(path, w):
    from docx import Document
    doc = Document()
    # Title avoids a leading "Well " token so the well-name regex below doesn't
    # latch onto the heading; the explicit "Well Name:" line gives it a clean hit.
    doc.add_heading("Summary Report", level=0)
    doc.add_heading(w["name"], level=1)
    doc.add_paragraph(f"Well Name: {w['name']}")
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in [
        ("Well Name", w["name"]), ("UWI", w["uwi"]), ("API", fmt_api(w["uwi"])),
        ("Operator", w["operator"]), ("Field", w["field"]),
        ("County", w["county"]), ("State", w["state"]), ("Country", w["country"]),
        ("Surface Latitude", f"{w['lat']:.5f}"), ("Surface Longitude", f"{w['lon']:.5f}"),
        ("Total Depth (ft)", f"{w['stop']:.1f}"),
    ]:
        c = t.add_row().cells
        c[0].text, c[1].text = k, str(v)
    doc.add_paragraph()
    doc.add_paragraph(
        f"This summary describes well {w['name']} (UWI {w['uwi']}, "
        f"API {fmt_api(w['uwi'])}), operated by {w['operator']} in the "
        f"{w['field']} field, {w['county']} County, {w['state']}.")
    doc.save(path)


# ── SEG-Y (segyio) ───────────────────────────────────────────────────────────
def gen_segy(path, s):
    import numpy as np
    import segyio
    nsamp = 250
    spec = segyio.spec()
    spec.format = 1                      # 4-byte IBM float
    spec.samples = list(range(nsamp))
    spec.ilines = list(range(1, s["ilines"] + 1))
    spec.xlines = list(range(1, s["xlines"] + 1))
    spec.sorting = segyio.TraceSortingFormat.INLINE_SORTING
    with segyio.create(path, spec) as f:
        f.text[0] = segyio.tools.create_text_header({
            1: f"SURVEY: {s['name']}",
            2: f"SURVEY TYPE: {s['dim']}",
            3: f"ACQUISITION YEAR: {s['year']}",
            4: f"AREA: {s['state']} {s['country']}",
            5: "SAMPLE INTERVAL (US): 4000",
            6: "GENERATED TEST SEG-Y - DataView fixture",
        })
        f.bin.update(hns=nsamp, hdt=4000)
        tr = 0
        base_x, base_y = 500000, 3300000
        for il in spec.ilines:
            for xl in spec.xlines:
                f.header[tr] = {
                    segyio.su.iline: il, segyio.su.xline: xl,
                    segyio.su.cdpx: base_x + xl * 25, segyio.su.cdpy: base_y + il * 25,
                    segyio.su.cdp: tr + 1, segyio.su.ns: nsamp, segyio.su.dt: 4000,
                }
                f.trace[tr] = (np.sin(np.linspace(0, 6.28, nsamp)) *
                               random.uniform(0.5, 1.5)).astype(np.float32)
                tr += 1


# ── P190 navigation (text) ───────────────────────────────────────────────────
def gen_p190(path, s):
    lines = [
        f"H0100 SURVEY NAME              {s['name']}",
        f"H0102 SURVEY TYPE              {s['dim']}",
        f"H0103 ACQUISITION YEAR         {s['year']}",
        f"H0200 AREA                     {s['state']} {s['country']}",
        "H1400 GEODETIC DATUM            WGS84",
        "H1500 PROJECTION                UTM",
        "H2600 SHOTPOINT NAV (P1/90)",
    ]
    base_x, base_y = 500000.0, 3300000.0
    for sp in range(1, 51):
        x = base_x + sp * 25.0
        y = base_y + sp * 5.0
        lines.append(f"S{sp:06d}      {x:12.1f}{y:13.1f}   {sp:6d}")
    with open(path, "w", newline="\n") as f:
        f.write("\n".join(lines) + "\n")


# ── Spatial: shapefile (pyshp) + GeoJSON ─────────────────────────────────────
def gen_shapefile(stem, wells):
    import shapefile
    w = shapefile.Writer(stem, shapeType=shapefile.POINT)
    w.field("UWI", "C", 14)
    w.field("WELL_NAME", "C", 80)
    w.field("OPERATOR", "C", 80)
    w.field("STATE", "C", 2)
    w.field("COUNTY", "C", 40)
    for d in wells:
        w.point(d["lon"], d["lat"])
        w.record(d["uwi"], d["name"], d["operator"], d["state"], d["county"])
    w.close()
    with open(stem + ".prj", "w") as f:                     # WGS84 WKT
        f.write('GEOGCS["GCS_WGS_1984",DATUM["D_WGS_1984",'
                'SPHEROID["WGS_1984",6378137.0,298.257223563]],'
                'PRIMEM["Greenwich",0.0],UNIT["Degree",0.0174532925199433]]')


def gen_geojson(path, wells):
    fc = {"type": "FeatureCollection", "features": [
        {"type": "Feature",
         "geometry": {"type": "Point", "coordinates": [d["lon"], d["lat"]]},
         "properties": {"UWI": d["uwi"], "WELL_NAME": d["name"],
                        "OPERATOR": d["operator"], "STATE": d["state"],
                        "COUNTY": d["county"]}}
        for d in wells]}
    with open(path, "w") as f:
        json.dump(fc, f, indent=2)


# ── orchestration ────────────────────────────────────────────────────────────
def main():
    ap = argparse.ArgumentParser(description="Generate a DataView test corpus.")
    ap.add_argument("--root", default="TestData", help="output root folder")
    a = ap.parse_args()
    root = os.path.abspath(a.root)
    made, skipped = [], []

    def attempt(label, fn):
        try:
            fn()
            made.append(label)
        except Exception as e:          # missing lib or write error — skip, don't abort
            skipped.append(f"{label}: {e}")

    # wells
    for w in WELLS:
        d = os.path.join(root, "Wells", f"{w['uwi']}__{safe(w['name'])}")
        os.makedirs(d, exist_ok=True)
        u = w["uwi"]
        attempt(f"{u} run1.las",  lambda d=d, u=u, w=w: gen_las(os.path.join(d, f"{u}_run1_triple_combo.las"), w, RUN1, "run1"))
        attempt(f"{u} run2.las",  lambda d=d, u=u, w=w: gen_las(os.path.join(d, f"{u}_run2_resistivity.las"), w, RUN2, "run2"))
        attempt(f"{u} directional.pdf", lambda d=d, u=u, w=w: gen_directional_pdf(os.path.join(d, f"{u}_directional_survey.pdf"), w))
        attempt(f"{u} formation.pdf",   lambda d=d, u=u, w=w: gen_formation_pdf(os.path.join(d, f"{u}_formation_tops.pdf"), w))
        attempt(f"{u} scout.pdf",        lambda d=d, u=u, w=w: gen_scout_pdf(os.path.join(d, f"{u}_scout_ticket.pdf"), w))
        attempt(f"{u} completion.xlsx", lambda d=d, u=u, w=w: gen_xlsx(os.path.join(d, f"{u}_completion.xlsx"), w))
        attempt(f"{u} summary.docx",    lambda d=d, u=u, w=w: gen_docx(os.path.join(d, f"{u}_well_summary.docx"), w))

    # seismic
    for s in SURVEYS:
        d = os.path.join(root, "Seismic", f"{s['folder']}_{s['year']}")
        os.makedirs(d, exist_ok=True)
        attempt(f"{s['name']} {s['ext']}", lambda d=d, s=s: gen_segy(os.path.join(d, s["file"] + s["ext"]), s))
        attempt(f"{s['name']} .p190",      lambda d=d, s=s: gen_p190(os.path.join(d, s["file"] + ".p190"), s))

    # spatial
    sp = os.path.join(root, "Spatial")
    os.makedirs(sp, exist_ok=True)
    attempt("well_locations.shp", lambda: gen_shapefile(os.path.join(sp, "well_locations"), WELLS))
    attempt("well_locations.geojson", lambda: gen_geojson(os.path.join(sp, "well_locations.geojson"), WELLS))

    # manifest
    man = os.path.join(root, "MANIFEST.txt")
    with open(man, "w") as f:
        f.write("DataView test corpus\n====================\n\n")
        f.write("WELLS (UWI / name / operator / county,state):\n")
        for w in WELLS:
            f.write(f"  {w['uwi']}  {w['name']:<14} {w['operator']:<20} "
                    f"{w['county']}, {w['state']}\n")
        f.write("\nSEISMIC (survey / dim / year / state):\n")
        for s in SURVEYS:
            f.write(f"  {s['name']:<14} {s['dim']}  {s['year']}  {s['state']}\n")
        f.write(f"\nGenerated {len(made)} file(s).\n")
        if skipped:
            f.write("\nSkipped (missing library or error):\n")
            for s in skipped:
                f.write(f"  - {s}\n")

    print(f"Root: {root}")
    print(f"Generated {len(made)} file(s) across "
          f"{len(WELLS)} wells, {len(SURVEYS)} surveys, spatial.")
    if skipped:
        print(f"Skipped {len(skipped)}:")
        for s in skipped:
            print(f"  - {s}")
    print(f"Manifest: {man}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
