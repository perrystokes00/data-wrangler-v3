"""
docx_document_loader.py — extract well data out of Word documents (final well reports,
completion reports, geological summaries) into bulk-loader staging CSVs.

Mirrors pdf_document_loader's output shape and API, so the same scan → review → map →
promote path handles it with no downstream change:

    extract_file(path, source="DOCX") -> {"doc_type", "uwi", "well_name",
                                          "well": [...], "formation": [...], ...}
    write_staging_csvs(directory, out_dir, source, files) -> {kind: (csv_path, nrows)}
    TARGET  -> {kind: DV_ table name}

A Word report is read as (text, tables) in pdfplumber's shape:
  * **spec tables** (Label | Value | Label | Value) are folded into the text as
    "Label: value" lines, so header fields map like a PDF's free-text header block;
  * **data tables** (header row + rows) stay as grids for rows_of()/_find_col().

Handles .docx (and .doc / .odt if LibreOffice is available to convert). Requires python-docx.
"""
from __future__ import annotations
import os
import re
import csv
import glob
import subprocess
import tempfile

# ── shared helpers (same semantics as pdf_document_loader) ───────────────────
def _num(s):
    """Digits/decimal/sign only — strips units, commas, degree signs."""
    s = str(s or "").strip()
    if not s or s in ("—", "-", "N/A", "None"):
        return ""
    m = re.search(r"-?\d[\d,]*\.?\d*", s)
    return m.group(0).replace(",", "") if m else ""


def _norm(s):
    return re.sub(r"[^a-z0-9]+", " ", str(s or "").lower()).strip()


def _cell(row, i):
    return str(row[i]).strip() if (i is not None and i < len(row) and row[i] is not None) else ""


def _find_col(head, *names):
    """Index of the first column whose normalized name contains any of `names`."""
    hn = [_norm(h) for h in head]
    for n in names:
        n = _norm(n)
        for i, h in enumerate(hn):
            if h == n:
                return i
        for i, h in enumerate(hn):
            if n and n in h:
                return i
    return None


# ── header field mapping (spec-table labels → dv_well columns) ───────────────
_HDR = {
    "api": "UWI", "api uwi": "UWI", "uwi": "UWI", "well id": "UWI",
    "well name": "WELL_NAME", "well": "WELL_NAME",
    "operator": "OPERATOR", "licensee": "LICENSEE",
    "well class": "WELL_CLASS", "well type": "WELL_CLASS",
    "status": "STATUS", "field": "FIELD_NAME", "field name": "FIELD_NAME",
    "formation at td": "FORMATION_AT_TD",
    "county": "COUNTY", "parish": "COUNTY",
    "state": "PROVINCE_STATE", "province": "PROVINCE_STATE",
    "state country": "PROVINCE_STATE", "country": "COUNTRY",
    "surface latitude": "SURFACE_LATITUDE", "latitude": "SURFACE_LATITUDE",
    "surface longitude": "SURFACE_LONGITUDE", "longitude": "SURFACE_LONGITUDE",
    "spud date": "SPUD_DATE", "completion date": "COMPLETION_DATE",
    "drillers td": "DRILLERS_TD", "driller s td": "DRILLERS_TD",
    "total depth": "DRILLERS_TD", "total depth md": "DRILLERS_TD",
    "depth datum": "DEPTH_DATUM", "kb elevation": "KB_ELEV", "gl elevation": "GL_ELEV",
}
_NUMERIC = {"SURFACE_LATITUDE", "SURFACE_LONGITUDE", "DRILLERS_TD", "KB_ELEV", "GL_ELEV"}

# extra labels that mark a Label|Value spec block (not dv_well fields, but they tell
# _is_spec_table that the table has no header row — e.g. the logging summary block)
_SPEC_EXTRA = {
    "log date", "log type", "top depth", "base depth", "interval logged", "source",
    "run", "run number", "core id", "core type", "recovery", "recovery pct",
    "report date", "prepared by", "rig", "rig name", "td date",
}

_DOC_HINTS = [
    ("eow", r"final well report|end of well|eow report|well completion report"),
    ("geo", r"geological (summary|report)|formation evaluation"),
    ("scout", r"scout ticket"),
]


def _detect_type(text):
    t = text.lower()
    for name, pat in _DOC_HINTS:
        if re.search(pat, t):
            return name
    return "unknown"


# ── read a Word file as (text, tables) ──────────────────────────────────────
def _is_spec_table(grid):
    """Label|Value|Label|Value — no header row. True when >=half the even columns
    of the first rows look like known header labels."""
    if not grid or len(grid[0]) < 2:
        return False
    hits = tot = 0
    for row in grid[:4]:
        for i in range(0, len(row) - 1, 2):
            lab = _norm(row[i])
            if lab:
                tot += 1
                if lab in _HDR or lab in _SPEC_EXTRA:
                    hits += 1
    return tot > 0 and hits >= max(1, tot // 2)


def _convert_to_docx(path):
    """.doc/.odt → .docx via LibreOffice, if available. Returns a temp path or None."""
    try:
        tmp = tempfile.mkdtemp(prefix="dv_docx_")
        subprocess.run(["soffice", "--headless", "--convert-to", "docx", "--outdir", tmp, path],
                       check=True, capture_output=True, timeout=120)
        out = glob.glob(os.path.join(tmp, "*.docx"))
        return out[0] if out else None
    except Exception:
        return None


def _read_docx(path):
    """→ (text, tables). Spec tables are folded into text as 'Label: value' lines."""
    import docx
    p = path
    if not path.lower().endswith(".docx"):
        p = _convert_to_docx(path)
        if not p:
            raise RuntimeError("cannot convert to .docx (LibreOffice unavailable)")
    d = docx.Document(p)
    lines = [x.text.strip() for x in d.paragraphs if x.text.strip()]
    spec, tables = [], []
    for t in d.tables:
        grid = []
        for r in t.rows:
            seen, row = set(), []
            for c in r.cells:                    # merged cells repeat — de-dup by identity
                if id(c._tc) in seen:
                    continue
                seen.add(id(c._tc))
                row.append(c.text.strip())
            grid.append(row)
        if not grid:
            continue
        if _is_spec_table(grid):
            for row in grid:
                for i in range(0, len(row) - 1, 2):
                    lab, val = row[i].strip(), row[i + 1].strip()
                    if lab and val:
                        spec.append(f"{lab}: {val}")
        else:
            tables.append(grid)
    return "\n".join(spec + lines), tables


def _header(text):
    """'Label: value' lines → {DV_WELL column: value}."""
    out = {}
    for line in text.splitlines():
        m = re.match(r"\s*([A-Za-z][A-Za-z /'’\.]{1,28}?)\s*[:\t]\s*(.+?)\s*$", line)
        if not m:
            continue
        raw_lab = _norm(m.group(1))
        v = m.group(2).strip()
        # combined "State / Country" (or "State / Province") → fill both sides
        if raw_lab in ("state country", "province country") and "/" in v:
            a, b = [x.strip() for x in v.split("/", 1)]
            out.setdefault("PROVINCE_STATE", a)
            out.setdefault("COUNTRY", b)
            continue
        key = _HDR.get(raw_lab)
        if key and key not in out:
            out[key] = _num(v) if key in _NUMERIC else v
    return out


# ── extract ─────────────────────────────────────────────────────────────────
def _depth_unit(text, tables):
    """The document states its own depth unit in headings/column headers ("Top MD (ft)",
    "Depth (m)"). Read it rather than assume; default FT (US land data) if truly absent."""
    blob = text + " " + " ".join(" ".join(str(c) for c in row)
                                 for t in tables for row in t[:1])
    b = blob.lower()
    if re.search(r"\(\s*m\s*\)|\bmetres?\b|\bmeters?\b", b) and not re.search(r"\(\s*ft\s*\)", b):
        return "M"
    return "FT"


def extract_file(path, source="DOCX"):
    res = {k: [] for k in ("well", "formation", "log", "curve", "core", "srvy_hdr", "srvy_sta")}
    res.update(doc_type="unknown", uwi="", well_name="", file=os.path.basename(path), error=None)
    try:
        text, tables = _read_docx(path)
    except Exception as e:
        res["error"] = str(e)
        return res

    res["doc_type"] = _detect_type(text)
    ouom = _depth_unit(text, tables)
    h = _header(text)
    uwi = _num(h.get("UWI", ""))[:14]
    wn = h.get("WELL_NAME", "")
    res["uwi"], res["well_name"] = uwi, wn
    stem = os.path.splitext(os.path.basename(path))[0]

    def rows_of(*sig):
        """First table whose header row contains all of `sig` → (head, body)."""
        for t in tables:
            if not t or not t[0]:
                continue
            hn = " ".join(_norm(c) for c in t[0])
            if all(_norm(s) in hn for s in sig):
                return t[0], t[1:]
        return None, []

    # ---- dv_well ----
    if uwi or wn:
        res["well"].append({
            "uwi": uwi, "well_name": wn, "operator_name": h.get("OPERATOR", ""),
            "well_type": h.get("WELL_CLASS", ""),
            "well_status": re.sub(r"\s*\(.*\)$", "", h.get("STATUS", "")).strip(),
            "field_name": h.get("FIELD_NAME", ""),
            "formation_at_td": h.get("FORMATION_AT_TD", ""),
            "country": h.get("COUNTRY", ""), "province_state": h.get("PROVINCE_STATE", ""),
            "county": h.get("COUNTY", ""),
            "surface_latitude": h.get("SURFACE_LATITUDE", ""),
            "surface_longitude": h.get("SURFACE_LONGITUDE", ""),
            "spud_date": h.get("SPUD_DATE", ""), "completion_date": h.get("COMPLETION_DATE", ""),
            "final_td": h.get("DRILLERS_TD", ""), "depth_datum": h.get("DEPTH_DATUM", ""),
            "kb_elevation": h.get("KB_ELEV", ""), "ground_elevation": h.get("GL_ELEV", ""),
            "elevation_ouom": ouom, "source": source})

    # ---- formation tops ----
    fh, fb = rows_of("top md")
    if fh:
        i_u = _find_col(fh, "strat unit", "formation", "unit")
        i_s = _find_col(fh, "strat name set", "name set")
        i_t = _find_col(fh, "top md", "top")
        i_b = _find_col(fh, "base md", "base")
        i_d = _find_col(fh, "interp date", "date")
        i_y = _find_col(fh, "interp by", "by")
        for r in fb:
            unit = _cell(r, i_u)
            if unit and _num(_cell(r, i_t)):
                res["formation"].append({
                    "uwi": uwi, "strat_unit_id": unit, "strat_unit_name": unit,
                    "strat_name_set": _cell(r, i_s), "interp_id": "1",
                    "top_depth": _num(_cell(r, i_t)), "base_depth": _num(_cell(r, i_b)),
                    "interp_date": _cell(r, i_d), "interpreter_ba_id": _cell(r, i_y),
                    "depth_ouom": ouom, "source": source})

    # ---- log header (from the spec block) + curves ----
    log_id = ""
    m = re.search(r"\b(LOG_[A-Z0-9_]+)\b", text)
    if m:
        log_id = m.group(1)
    lt = re.search(r"\b(WIRELINE|LWD|MWD|MUD LOG)\b", text.upper())
    ld = re.search(r"Log Date:\s*([\d/-]+)", text)
    top = re.search(r"Top Depth:\s*([\d,\.]+)", text)
    bas = re.search(r"Base Depth:\s*([\d,\.]+)", text)
    run = re.search(r"Run\s*(\d+)", text)
    if log_id or lt:
        res["log"].append({
            "uwi": uwi, "log_id": log_id or f"LOG_{uwi}", "log_type": lt.group(1) if lt else "",
            "run_num": run.group(1) if run else "", "log_date": ld.group(1) if ld else "",
            "top_depth": _num(top.group(1)) if top else "",
            "base_depth": _num(bas.group(1)) if bas else "", "depth_ouom": ouom, "source": source})

    ch, cb = rows_of("curve", "unit")
    if ch:
        i_c = _find_col(ch, "curve")
        i_u = _find_col(ch, "unit")
        i_mn = _find_col(ch, "min value", "min")
        i_mx = _find_col(ch, "max value", "max")
        lid = log_id or f"LOG_{uwi}"
        seen_cv = {}
        for r in cb:
            name = _cell(r, i_c)
            if name:
                # curve_id is NOT NULL and has no source column — generate it here rather
                # than make the operator write a concat rule. log_id already carries the
                # uwi, so {log_id}_{mnemonic} is unique and fits nvarchar(40). A repeated
                # mnemonic in one log gets _2, _3 so the PK can't collide.
                n = seen_cv.get(name, 0) + 1
                seen_cv[name] = n
                cid = f"{lid}_{name}" if n == 1 else f"{lid}_{name}_{n}"
                res["curve"].append({
                    "uwi": uwi, "log_id": lid, "curve_id": cid[:40], "mnemonic": name,
                    "curve_description": "", "curve_unit": _cell(r, i_u),
                    "min_value": _num(_cell(r, i_mn)),
                    "max_value": _num(_cell(r, i_mx)), "depth_ouom": ouom, "source": source})

    # ---- core ----
    coh, cob = rows_of("core id")
    if not coh:
        coh, cob = rows_of("recovery")
    if coh:
        i_id = _find_col(coh, "core id")
        i_ty = _find_col(coh, "type")
        i_t = _find_col(coh, "top")
        i_b = _find_col(coh, "base")
        i_r = _find_col(coh, "recovery")
        i_f = _find_col(coh, "formation")
        for k, r in enumerate(cob, 1):
            t_ = _num(_cell(r, i_t))
            if not t_:
                continue
            res["core"].append({
                "uwi": uwi, "core_id": _cell(r, i_id) or f"CORE_{uwi}_{k}",
                "core_type": _cell(r, i_ty), "top_depth": t_, "base_depth": _num(_cell(r, i_b)),
                "recovery_pct": _num(_cell(r, i_r)), "strat_unit_name": _cell(r, i_f),
                "depth_ouom": ouom, "length_ouom": ouom, "source": source})

    # ---- directional survey ----
    sh, sb = rows_of("md", "azimuth")
    if not sh:
        sh, sb = rows_of("md", "azi")
    if sh:
        srvy_id = f"{uwi}-SRVY" if uwi else f"SRVY_{stem}"
        i_q = _find_col(sh, "seq")
        i_m = _find_col(sh, "md")
        i_i = _find_col(sh, "inclination", "inc")
        i_a = _find_col(sh, "azimuth", "azi")
        i_v = _find_col(sh, "tvdss", "tvd")
        sta = []
        for k, r in enumerate(sb, 1):
            md = _num(_cell(r, i_m))
            if not md:
                continue
            try:                                   # real stations only
                inc = float(_num(_cell(r, i_i)) or 999)
                azi = float(_num(_cell(r, i_a)) or 999)
            except ValueError:
                continue
            if inc > 120 or azi > 360:
                continue
            sta.append({"uwi": uwi, "survey_id": srvy_id,
                        "station_id": _cell(r, i_q) or str(k), "md": md,
                        "incl": _num(_cell(r, i_i)), "azim": _num(_cell(r, i_a)),
                        "tvd": _num(_cell(r, i_v)), "depth_ouom": ouom, "source": source})
        if sta:
            res["srvy_hdr"].append({"uwi": uwi, "survey_id": srvy_id, "source": source})
            res["srvy_sta"] = sta
    return res


# ── staging CSVs ────────────────────────────────────────────────────────────
# Column names are the TARGET TABLE's real column names (from INFORMATION_SCHEMA), so the
# loader's exact-name matcher maps every one automatically — no hand-mapping, no function
# rules for keys. Don't rename these to suit a source document; the schema is the contract.
_COLS = {
    "well": ["uwi", "well_name", "operator_name", "well_type", "well_status", "field_name",
             "formation_at_td", "country", "province_state", "county", "surface_latitude",
             "surface_longitude", "spud_date", "completion_date", "final_td", "depth_datum",
             "kb_elevation", "ground_elevation", "elevation_ouom", "source"],
    "formation": ["uwi", "strat_unit_id", "strat_unit_name", "strat_name_set", "interp_id",
                  "top_depth", "base_depth", "interp_date", "interpreter_ba_id", "depth_ouom",
                  "source"],
    "log": ["uwi", "log_id", "log_type", "run_num", "log_date", "top_depth", "base_depth",
            "depth_ouom", "source"],
    "curve": ["uwi", "log_id", "curve_id", "mnemonic", "curve_description", "curve_unit",
              "min_value", "max_value", "depth_ouom", "source"],
    "core": ["uwi", "core_id", "core_type", "top_depth", "base_depth", "recovery_pct",
             "strat_unit_name", "depth_ouom", "length_ouom", "source"],
    "srvy_hdr": ["uwi", "survey_id", "source"],
    "srvy_sta": ["uwi", "survey_id", "station_id", "md", "incl", "azim", "tvd",
                 "depth_ouom", "source"],
}

TARGET = {
    "well": "DV_WELL", "formation": "DV_WELL_FORMATION_TOP", "log": "DV_WELL_LOG",
    "curve": "DV_WELL_LOG_CURVE", "core": "DV_WELL_CORE",
    "srvy_hdr": "DV_WELL_DIR_SRVY_HDR", "srvy_sta": "DV_WELL_DIR_SRVY_STA",
}

EXTS = (".docx", ".doc", ".odt")


def write_staging_csvs(directory, out_dir=None, source="DOCX", files=None):
    """Extract every Word doc in `directory` → one CSV per kind. {kind: (path, nrows)}."""
    out_dir = out_dir or directory
    os.makedirs(out_dir, exist_ok=True)
    if files is None:
        files = []
        for e in EXTS:
            files += glob.glob(os.path.join(directory, f"*{e}"))
            files += glob.glob(os.path.join(directory, f"*{e.upper()}"))
        files = sorted(f for f in files if not os.path.basename(f).startswith("~$"))
    agg = {k: [] for k in _COLS}
    for p in files:
        r = extract_file(p, source=source)
        if r.get("error"):
            continue
        for k in _COLS:
            agg[k].extend(r.get(k) or [])
    written = {}
    for k, rowset in agg.items():
        if not rowset:
            continue
        path = os.path.join(out_dir, f"docx_{k}.csv")
        with open(path, "w", newline="", encoding="utf-8") as fh:
            w = csv.DictWriter(fh, fieldnames=_COLS[k], extrasaction="ignore")
            w.writeheader()
            w.writerows(rowset)
        written[k] = (path, len(rowset))
    return written


if __name__ == "__main__":
    import sys
    d = sys.argv[1] if len(sys.argv) > 1 else "."
    for k, (p, n) in write_staging_csvs(d).items():
        print(f"{k:12} {n:4} rows -> {p}")
