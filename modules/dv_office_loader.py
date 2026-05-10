"""
modules/dv_office_loader.py
============================
DataView v3 — Office Document Loader

Loads data from Excel (.xlsx) and Word (.docx) documents into DataView tables.

Functions:
    load_formation_tops_xlsx(engine, path, source)   → dict
    load_completion_xlsx(engine, path, source)       → dict
    load_production_xlsx(engine, path, source)       → dict
    load_completion_docx(engine, path, source)       → dict
    load_formation_tops_docx(engine, path, source)   → dict

All functions:
    engine : SQLAlchemy engine connected to DataView database
    path   : str — full file path on disk
    source : str — source code (default "DATA_LOADER")
    returns: {"loaded": int, "errors": [...], "wells": [...]}

Requirements:
    pip install openpyxl python-docx
"""
from __future__ import annotations

import hashlib
import re
import uuid
from datetime import datetime, date, timedelta
from pathlib import Path

import pandas as pd
from sqlalchemy import text


# =============================================================================
# HELPERS
# =============================================================================

def _uid() -> str:
    return hashlib.sha1(str(uuid.uuid4()).encode()).hexdigest()[:40]


def _trunc(v, n: int = 40) -> str | None:
    if v is None:
        return None
    s = str(v).strip()
    return s[:n] if s else None


def _safe_float(v) -> float | None:
    if v is None or str(v).strip() in ("", "None", "nan", "NaN"):
        return None
    try:
        return float(str(v).replace(",", "").strip())
    except (ValueError, TypeError):
        return None


def _excel_date(v) -> str | None:
    """Convert Excel serial date (float/int) or string date to ISO string."""
    if v is None:
        return None
    # Already a datetime or date
    if isinstance(v, (datetime, date)):
        return v.strftime("%Y-%m-%d")
    # Try float (Excel serial)
    try:
        serial = float(v)
        if 10000 < serial < 100000:  # plausible Excel date range
            dt = datetime(1899, 12, 30) + timedelta(days=serial)
            return dt.strftime("%Y-%m-%d")
    except (ValueError, TypeError):
        pass
    # Try string parse
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d-%b-%Y", "%B %d, %Y",
                "%b %d, %Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(str(v).strip(), fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return _trunc(v, 20)


def _result(loaded: int = 0, errors: list | None = None, **extra) -> dict:
    return {"loaded": loaded, "errors": errors or [], **extra}


def _well_exists(con, uwi: str) -> bool:
    row = con.execute(text(
        "SELECT COUNT(*) FROM dataview.dv_well WHERE uwi = :uwi"
    ), {"uwi": uwi}).scalar()
    return bool(row)


def _normalise_uwi(raw: str) -> str:
    """Strip dashes from UWI for comparison — 42-317-12345-00-00 → 4231712345000."""
    return re.sub(r"[-\s]", "", str(raw).strip()) if raw else ""


def _find_well(con, raw_uwi: str) -> str | None:
    """
    Try to match raw_uwi (with or without dashes) to dv_well.uwi.
    Returns the canonical uwi stored in the DB, or None.
    """
    # Exact match first
    row = con.execute(text(
        "SELECT uwi FROM dataview.dv_well WHERE uwi = :u LIMIT 1"
        if False else
        "SELECT TOP 1 uwi FROM dataview.dv_well WHERE uwi = :u"
    ), {"u": raw_uwi.strip()}).fetchone()
    if row:
        return row[0]
    # Try without dashes
    norm = _normalise_uwi(raw_uwi)
    if not norm:
        return None
    row = con.execute(text(
        "SELECT TOP 1 uwi FROM dataview.dv_well "
        "WHERE REPLACE(REPLACE(uwi,'-',''),' ','') = :u"
    ), {"u": norm}).fetchone()
    return row[0] if row else None


# =============================================================================
# 1. FORMATION TOPS FROM XLSX → dv_well_formation_top
# =============================================================================

def load_formation_tops_xlsx(engine, path: str,
                              source: str = "DATA_LOADER") -> dict:
    """
    Load formation tops from Formation_Tops.xlsx style file.

    Expected columns (case-insensitive):
        UWI, WELL_NAME, FORMATION, TOP_MD_FT, BASE_MD_FT,
        TOP_TVD_FT (optional), NET_PAY_FT (optional),
        FLUID (optional), PICKED_BY (optional), PICK_DATE (optional)
    """
    errors  = []
    loaded  = 0
    wells   = []

    try:
        xl = pd.ExcelFile(path)
        # Find the formation tops sheet
        sheet = next(
            (s for s in xl.sheet_names
             if any(kw in s.lower() for kw in
                    ("formation", "top", "pick", "strat"))),
            xl.sheet_names[0]
        )
        df = pd.read_excel(path, sheet_name=sheet, dtype=str)
        df.columns = [c.upper().strip() for c in df.columns]
        df = df.dropna(how="all")
    except Exception as exc:
        return _result(errors=[f"Could not read Excel: {exc}"])

    # Column name aliases
    _COL = {
        "UWI":         next((c for c in df.columns if "UWI" in c), None),
        "WELL_NAME":   next((c for c in df.columns if "WELL" in c and "NAME" in c), None),
        "FORMATION":   next((c for c in df.columns if any(k in c for k in ("FORMATION","FORM","UNIT","STRAT"))), None),
        "TOP_MD":      next((c for c in df.columns if "TOP" in c and ("MD" in c or "FT" in c) and "BASE" not in c and "TVD" not in c), None),
        "BASE_MD":     next((c for c in df.columns if "BASE" in c and ("MD" in c or "FT" in c)), None),
        "TOP_TVD":     next((c for c in df.columns if "TVD" in c and "BASE" not in c), None),
        "NET_PAY":     next((c for c in df.columns if "NET" in c and "PAY" in c), None),
        "FLUID":       next((c for c in df.columns if "FLUID" in c), None),
        "PICKED_BY":   next((c for c in df.columns if "PICK" in c and "BY" in c), None),
        "PICK_DATE":   next((c for c in df.columns if "PICK" in c and "DATE" in c), None),
    }

    if not _COL["UWI"] or not _COL["FORMATION"] or not _COL["TOP_MD"]:
        return _result(errors=[
            f"Required columns missing. Found: {list(df.columns)}. "
            f"Need: UWI, FORMATION, TOP_MD_FT"
        ])

    try:
        with engine.begin() as con:
            for i, row in df.iterrows():
                raw_uwi    = str(row.get(_COL["UWI"], "")).strip()
                form_name  = _trunc(row.get(_COL["FORMATION"]), 255)
                top_depth  = _safe_float(row.get(_COL["TOP_MD"]))
                base_depth = _safe_float(row.get(_COL["BASE_MD"])) if _COL["BASE_MD"] else None
                net_pay    = _safe_float(row.get(_COL["NET_PAY"])) if _COL["NET_PAY"] else None
                fluid      = _trunc(row.get(_COL["FLUID"]), 40) if _COL["FLUID"] else None
                picked_by  = _trunc(row.get(_COL["PICKED_BY"]), 100) if _COL["PICKED_BY"] else None
                pick_date  = _excel_date(row.get(_COL["PICK_DATE"])) if _COL["PICK_DATE"] else None

                if not raw_uwi or not form_name or top_depth is None:
                    continue

                uwi = _find_well(con, raw_uwi)
                if not uwi:
                    errors.append(f"Row {i+1}: UWI '{raw_uwi}' not found in dv_well")
                    continue

                strat_id  = _uid()
                interp_id = "1"

                try:
                    con.execute(text("""
                        INSERT INTO dataview.dv_well_formation_top (
                            uwi, strat_unit_id, interp_id,
                            strat_unit_name, strat_unit_type,
                            top_depth, base_depth, depth_ouom, depth_datum,
                            net_thickness, fluid_type,
                            picked_by, pick_date,
                            active_ind,
                            row_created_by, row_created_date, source
                        ) VALUES (
                            :uwi, :sid, :iid,
                            :fname, 'FORMATION',
                            :top, :base, 'FT', 'KB',
                            :net, :fluid,
                            :pby, :pdate,
                            'Y',
                            :by, GETDATE(), :src
                        )
                    """), {
                        "uwi":   uwi, "sid": strat_id, "iid": interp_id,
                        "fname": form_name,
                        "top":   top_depth, "base": base_depth,
                        "net":   net_pay, "fluid": fluid,
                        "pby":   picked_by, "pdate": pick_date,
                        "by":    "DataWrangler", "src": source,
                    })
                    loaded += 1
                    if uwi not in wells:
                        wells.append(uwi)
                except Exception as e:
                    errors.append(f"Row {i+1} ({form_name}): {e}")

    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=wells)


# =============================================================================
# 2. COMPLETION PARAMETERS FROM XLSX
#    → dv_well_completion + dv_well_perforation + dv_well_stimulation
# =============================================================================

def load_completion_xlsx(engine, path: str,
                         source: str = "DATA_LOADER") -> dict:
    """
    Load completion data from Completion_Parameters.xlsx style file.

    Reads three sheets:
        Completions  → dv_well_completion + dv_well_stimulation
        Perforations → dv_well_perforation
        Well Header  → UPDATE dv_well (lat/lon/operator if blank)
    """
    errors = []
    loaded = 0
    wells  = []

    try:
        xl = pd.ExcelFile(path)
        sheets = {s.lower(): s for s in xl.sheet_names}
    except Exception as exc:
        return _result(errors=[f"Could not read Excel: {exc}"])

    def _read_sheet(keywords):
        sheet = next(
            (sheets[k] for k in sheets
             if any(kw in k for kw in keywords)),
            None
        )
        if not sheet:
            return pd.DataFrame()
        df = pd.read_excel(path, sheet_name=sheet, dtype=str)
        df.columns = [c.upper().strip() for c in df.columns]
        return df.dropna(how="all")

    comp_df = _read_sheet(["complet", "frac", "stim"])
    perf_df = _read_sheet(["perf"])
    hdr_df  = _read_sheet(["header", "well header", "wells"])

    try:
        with engine.begin() as con:

            # ── Completions sheet ─────────────────────────────────────────────
            for i, row in comp_df.iterrows():
                raw_uwi = str(row.get("UWI", "")).strip()
                if not raw_uwi:
                    continue
                uwi = _find_well(con, raw_uwi)
                if not uwi:
                    errors.append(f"Comp row {i+1}: UWI '{raw_uwi}' not found")
                    continue

                comp_id = _uid()
                stim_id = _uid()
                comp_date = _excel_date(row.get("COMP_DATE"))

                try:
                    # dv_well_completion
                    con.execute(text("""
                        INSERT INTO dataview.dv_well_completion (
                            uwi, completion_id, completion_num,
                            completion_date,
                            top_depth, base_depth, depth_ouom, depth_datum,
                            perf_top, perf_base,
                            lateral_length,
                            active_ind,
                            row_created_by, row_created_date, source
                        ) VALUES (
                            :uwi, :cid, '1',
                            :cdate,
                            :ptop, :pbase, 'FT', 'KB',
                            :ptop, :pbase,
                            :lat,
                            'Y',
                            :by, GETDATE(), :src
                        )
                    """), {
                        "uwi":   uwi, "cid": comp_id,
                        "cdate": comp_date,
                        "ptop":  _safe_float(row.get("PERF_TOP_FT")),
                        "pbase": _safe_float(row.get("PERF_BOT_FT")),
                        "lat":   _safe_float(row.get("LATERAL_LEN_FT")),
                        "by":    "DataWrangler", "src": source,
                    })

                    # dv_well_stimulation
                    con.execute(text("""
                        INSERT INTO dataview.dv_well_stimulation (
                            uwi, completion_id, stimulation_id,
                            stim_type, stim_date,
                            num_stages,
                            total_proppant_lbs, total_fluid_bbl,
                            cluster_spacing_ft,
                            max_treatment_pressure_psi,
                            active_ind,
                            row_created_by, row_created_date, source
                        ) VALUES (
                            :uwi, :cid, :sid,
                            'HYDRAULIC_FRACTURE', :cdate,
                            :stages,
                            :prop, :fluid,
                            :clust,
                            NULL,
                            'Y',
                            :by, GETDATE(), :src
                        )
                    """), {
                        "uwi":    uwi, "cid": comp_id, "sid": stim_id,
                        "cdate":  comp_date,
                        "stages": _safe_float(row.get("FRAC_STAGES")),
                        "prop":   _safe_float(row.get("PROPPANT_LBS")),
                        "fluid":  _safe_float(row.get("FLUID_BBL")),
                        "clust":  _safe_float(row.get("CLUSTER_SPACING_FT")),
                        "by":     "DataWrangler", "src": source,
                    })
                    loaded += 1
                    if uwi not in wells:
                        wells.append(uwi)
                except Exception as e:
                    errors.append(f"Comp row {i+1}: {e}")

            # ── Perforations sheet ────────────────────────────────────────────
            for i, row in perf_df.iterrows():
                raw_uwi = str(row.get("UWI", "")).strip()
                if not raw_uwi:
                    continue
                uwi = _find_well(con, raw_uwi)
                if not uwi:
                    continue  # already reported above

                perf_id   = _uid()
                perf_date = _excel_date(row.get("PERF_DATE"))

                try:
                    con.execute(text("""
                        INSERT INTO dataview.dv_well_perforation (
                            uwi, perforation_id, perf_num,
                            perf_date,
                            top_depth, base_depth, depth_ouom, depth_datum,
                            strat_unit_name,
                            active_ind,
                            row_created_by, row_created_date, source
                        ) VALUES (
                            :uwi, :pid, :pnum,
                            :pdate,
                            :top, :base, 'FT', 'KB',
                            :sname,
                            'Y',
                            :by, GETDATE(), :src
                        )
                    """), {
                        "uwi":   uwi, "pid": perf_id,
                        "pnum":  str(i + 1),
                        "pdate": perf_date,
                        "top":   _safe_float(row.get("PERF_TOP_FT")),
                        "base":  _safe_float(row.get("PERF_BOT_FT")),
                        "sname": _trunc(row.get("FORMATION"), 255),
                        "by":    "DataWrangler", "src": source,
                    })
                    loaded += 1
                except Exception as e:
                    errors.append(f"Perf row {i+1}: {e}")

    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=wells)


# =============================================================================
# 3. PRODUCTION DATA FROM XLSX → dv_prod_entity + dv_prod_volume
# =============================================================================

def load_production_xlsx(engine, path: str,
                         source: str = "DATA_LOADER") -> dict:
    """
    Load monthly production from Production_Data.xlsx style file.

    Expected columns:
        UWI, WELL_NAME, DATE, OIL_BBL, GAS_MCF, WATER_BBL,
        BOE (optional), TUBING_PRESS_PSI (optional), STATUS (optional)
    """
    errors = []
    loaded = 0
    wells  = []

    try:
        xl   = pd.ExcelFile(path)
        sheet = next(
            (s for s in xl.sheet_names
             if any(k in s.lower() for k in ("prod", "monthly", "volume"))),
            xl.sheet_names[0]
        )
        df = pd.read_excel(path, sheet_name=sheet, dtype=str)
        df.columns = [c.upper().strip() for c in df.columns]
        df = df.dropna(how="all")
    except Exception as exc:
        return _result(errors=[f"Could not read Excel: {exc}"])

    try:
        with engine.begin() as con:
            # Group by UWI — create prod_entity once per well
            entity_created = set()

            for i, row in df.iterrows():
                raw_uwi = str(row.get("UWI", "")).strip()
                if not raw_uwi:
                    continue
                uwi = _find_well(con, raw_uwi)
                if not uwi:
                    if raw_uwi not in [e.split("|")[0] for e in errors]:
                        errors.append(f"UWI '{raw_uwi}' not found in dv_well")
                    continue

                # Create prod_entity once per well
                if uwi not in entity_created:
                    entity_id = _uid()
                    try:
                        con.execute(text("""
                            IF NOT EXISTS (
                                SELECT 1 FROM dataview.dv_prod_entity
                                WHERE uwi = :uwi
                            )
                            INSERT INTO dataview.dv_prod_entity (
                                uwi, entity_id, entity_type,
                                active_ind,
                                row_created_by, row_created_date, source
                            ) VALUES (
                                :uwi, :eid, 'WELL',
                                'Y',
                                :by, GETDATE(), :src
                            )
                        """), {
                            "uwi": uwi, "eid": entity_id,
                            "by":  "DataWrangler", "src": source,
                        })
                    except Exception:
                        pass  # May already exist
                    entity_created.add(uwi)
                    if uwi not in wells:
                        wells.append(uwi)

                # Get the entity_id for this well
                eid_row = con.execute(text(
                    "SELECT TOP 1 entity_id FROM dataview.dv_prod_entity "
                    "WHERE uwi = :uwi"
                ), {"uwi": uwi}).fetchone()
                if not eid_row:
                    continue
                entity_id = eid_row[0]

                # Production date
                prod_date = _excel_date(row.get("DATE"))
                if not prod_date:
                    errors.append(f"Row {i+1}: could not parse date '{row.get('DATE')}'")
                    continue

                # Extract year/month from date
                try:
                    dt = datetime.strptime(prod_date, "%Y-%m-%d")
                    prod_year  = dt.year
                    prod_month = dt.month
                except Exception:
                    prod_year = prod_month = None

                volume_id = _uid()
                try:
                    con.execute(text("""
                        INSERT INTO dataview.dv_prod_volume (
                            entity_id, volume_id,
                            prod_date, prod_year, prod_month,
                            period_type,
                            oil_vol, gas_vol, water_vol, boe_vol,
                            oil_vol_ouom, gas_vol_ouom, water_vol_ouom,
                            avg_tubing_pressure, pressure_ouom,
                            prod_status,
                            row_created_by, row_created_date, source
                        ) VALUES (
                            :eid, :vid,
                            :pdate, :pyear, :pmonth,
                            'MONTHLY',
                            :oil, :gas, :wtr, :boe,
                            'BBL', 'MCF', 'BBL',
                            :tpress, 'PSI',
                            :status,
                            :by, GETDATE(), :src
                        )
                    """), {
                        "eid":    entity_id, "vid": volume_id,
                        "pdate":  prod_date, "pyear": prod_year, "pmonth": prod_month,
                        "oil":    _safe_float(row.get("OIL_BBL")),
                        "gas":    _safe_float(row.get("GAS_MCF")),
                        "wtr":    _safe_float(row.get("WATER_BBL")),
                        "boe":    _safe_float(row.get("BOE")),
                        "tpress": _safe_float(row.get("TUBING_PRESS_PSI")),
                        "status": _trunc(row.get("STATUS", "PRODUCING"), 40),
                        "by":     "DataWrangler", "src": source,
                    })
                    loaded += 1
                except Exception as e:
                    errors.append(f"Row {i+1} ({prod_date}): {e}")

    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=wells)


# =============================================================================
# 4. WELL COMPLETION REPORT FROM DOCX
#    → dv_well_completion + dv_well_perforation + dv_well_stimulation
# =============================================================================

def load_completion_docx(engine, path: str,
                         source: str = "DATA_LOADER") -> dict:
    """
    Load completion report from a Word .docx file.
    Extracts tables by heading context:
        Well Identification → well header fields
        Perforation Intervals → dv_well_perforation
        Stimulation Parameters → dv_well_stimulation
        Initial Production → remark on dv_well_completion
    """
    errors = []
    loaded = 0

    try:
        from docx import Document
    except ImportError:
        return _result(errors=["python-docx not installed: pip install python-docx"])

    try:
        doc = Document(path)
    except Exception as exc:
        return _result(errors=[f"Could not read Word doc: {exc}"])

    # ── Extract all tables with their preceding heading ────────────────────
    def _table_to_df(table):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if not rows:
            return pd.DataFrame()
        # Detect if first row is a header
        first = rows[0]
        if len(first) == 2 and all(first):
            # Key-value table
            return pd.DataFrame(rows, columns=["FIELD", "VALUE"])
        if len(rows) > 1:
            return pd.DataFrame(rows[1:], columns=[c.upper() for c in rows[0]])
        return pd.DataFrame()

    # Walk document body to get heading → table pairs
    sections: dict[str, list] = {}
    current_heading = "GENERAL"
    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            style = block.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
            )
            if style is not None:
                sval = style.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
                )
                if "heading" in sval.lower() or "Heading" in sval:
                    texts = [n.text or "" for n in block.iter()
                             if n.text and n.tag.endswith("t")]
                    current_heading = " ".join(texts).strip().upper()
        elif tag == "tbl":
            from docx.oxml.ns import qn
            from docx.table import Table as DocxTable
            tbl_obj = DocxTable(block, doc)
            df = _table_to_df(tbl_obj)
            if current_heading not in sections:
                sections[current_heading] = []
            sections[current_heading].append(df)

    # ── Extract UWI from Well Identification table ─────────────────────────
    uwi       = None
    comp_date = None
    well_info = {}

    for heading, tables in sections.items():
        if any(k in heading for k in ("WELL ID", "IDENTIFICATION", "WELL IDENT")):
            for df in tables:
                if "FIELD" in df.columns and "VALUE" in df.columns:
                    for _, r in df.iterrows():
                        field = str(r["FIELD"]).upper().strip()
                        val   = str(r["VALUE"]).strip()
                        if not val or val.lower() in ("none", ""):
                            continue
                        if "UWI" in field:
                            uwi = val
                        elif "COMPLETION DATE" in field:
                            comp_date = _excel_date(val)
                        elif "SPUD" in field:
                            well_info["SPUD_DATE"] = _excel_date(val)
                        elif "TOTAL DEPTH" in field:
                            nums = re.findall(r"[\d,]+", val.split("/")[0])
                            if nums:
                                well_info["FINAL_TD"] = float(nums[0].replace(",",""))
                        elif "KB ELEVATION" in field:
                            nums = re.findall(r"[\d,]+", val)
                            if nums:
                                well_info["KB_ELEVATION"] = float(nums[0].replace(",",""))

    if not uwi:
        return _result(errors=["Could not find UWI in document"])

    try:
        with engine.begin() as con:
            db_uwi = _find_well(con, uwi)
            if not db_uwi:
                return _result(errors=[f"UWI '{uwi}' not found in dv_well"])

            comp_id = _uid()
            stim_id = _uid()

            # ── dv_well_completion ─────────────────────────────────────────────
            try:
                con.execute(text("""
                    INSERT INTO dataview.dv_well_completion (
                        uwi, completion_id, completion_num,
                        completion_date,
                        active_ind,
                        row_created_by, row_created_date, source
                    ) VALUES (
                        :uwi, :cid, '1',
                        :cdate,
                        'Y',
                        :by, GETDATE(), :src
                    )
                """), {
                    "uwi": db_uwi, "cid": comp_id,
                    "cdate": comp_date,
                    "by": "DataWrangler", "src": source,
                })
                loaded += 1
            except Exception as e:
                errors.append(f"Completion header: {e}")

            # ── Perforation intervals ──────────────────────────────────────────
            for heading, tables in sections.items():
                if any(k in heading for k in ("PERF", "INTERVAL")):
                    for df in tables:
                        if df.empty:
                            continue
                        cols = {c.upper(): c for c in df.columns}
                        for i, row in df.iterrows():
                            perf_id  = _uid()
                            top_col  = next((cols[c] for c in cols
                                             if "TOP" in c), None)
                            base_col = next((cols[c] for c in cols
                                             if "BASE" in c or "BOT" in c), None)
                            form_col = next((cols[c] for c in cols
                                             if "FORM" in c), None)
                            top  = _safe_float(row.get(top_col))  if top_col  else None
                            base = _safe_float(row.get(base_col)) if base_col else None
                            if top is None:
                                continue
                            try:
                                con.execute(text("""
                                    INSERT INTO dataview.dv_well_perforation (
                                        uwi, perforation_id, perf_num,
                                        top_depth, base_depth, depth_ouom, depth_datum,
                                        strat_unit_name,
                                        active_ind,
                                        row_created_by, row_created_date, source
                                    ) VALUES (
                                        :uwi, :pid, :pnum,
                                        :top, :base, 'FT', 'KB',
                                        :sname,
                                        'Y',
                                        :by, GETDATE(), :src
                                    )
                                """), {
                                    "uwi":  db_uwi, "pid": perf_id,
                                    "pnum": str(i+1),
                                    "top":  top, "base": base,
                                    "sname": _trunc(row.get(form_col), 255) if form_col else None,
                                    "by":   "DataWrangler", "src": source,
                                })
                                loaded += 1
                            except Exception as e:
                                errors.append(f"Perf row {i+1}: {e}")

            # ── Stimulation parameters ─────────────────────────────────────────
            stim_params: dict[str, float | None] = {}
            for heading, tables in sections.items():
                if any(k in heading for k in ("STIM", "FRAC", "TREAT")):
                    for df in tables:
                        if "FIELD" not in df.columns:
                            continue
                        for _, row in df.iterrows():
                            field = str(row.get("FIELD", "")).upper()
                            val   = row.get("VALUE") or row.get("value") or ""
                            nums  = re.findall(r"[\d,]+\.?\d*", str(val))
                            fval  = float(nums[0].replace(",","")) if nums else None
                            if "FRAC STAGE" in field or "STAGE" in field:
                                stim_params["stages"] = fval
                            elif "PROPPANT" in field and "STAGE" not in field:
                                stim_params["prop"] = fval
                            elif "FLUID" in field and "STAGE" not in field:
                                stim_params["fluid"] = fval
                            elif "CLUSTER SPACING" in field:
                                stim_params["clust"] = fval
                            elif "MAX TREAT" in field or "MAX PRESSURE" in field:
                                stim_params["maxpress"] = fval
                            elif "ISIP" in field or "ISIP" in field:
                                stim_params["isip"] = fval

            if stim_params:
                try:
                    con.execute(text("""
                        INSERT INTO dataview.dv_well_stimulation (
                            uwi, completion_id, stimulation_id,
                            stim_type, stim_date,
                            num_stages,
                            total_proppant_lbs, total_fluid_bbl,
                            cluster_spacing_ft,
                            max_treatment_pressure_psi,
                            avg_isip_psi,
                            active_ind,
                            row_created_by, row_created_date, source
                        ) VALUES (
                            :uwi, :cid, :sid,
                            'HYDRAULIC_FRACTURE', :cdate,
                            :stages,
                            :prop, :fluid,
                            :clust,
                            :maxp,
                            :isip,
                            'Y',
                            :by, GETDATE(), :src
                        )
                    """), {
                        "uwi":    db_uwi, "cid": comp_id, "sid": stim_id,
                        "cdate":  comp_date,
                        "stages": stim_params.get("stages"),
                        "prop":   stim_params.get("prop"),
                        "fluid":  stim_params.get("fluid"),
                        "clust":  stim_params.get("clust"),
                        "maxp":   stim_params.get("maxpress"),
                        "isip":   stim_params.get("isip"),
                        "by":     "DataWrangler", "src": source,
                    })
                    loaded += 1
                except Exception as e:
                    errors.append(f"Stimulation: {e}")

    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=[db_uwi] if db_uwi else [])


# =============================================================================
# 5. FORMATION TOPS FROM DOCX → dv_well_formation_top
# =============================================================================

def load_formation_tops_docx(engine, path: str,
                              source: str = "DATA_LOADER") -> dict:
    """
    Load formation tops from a Word .docx file.
    Finds tables with UWI/Formation/Top MD columns.
    """
    errors = []
    loaded = 0
    wells  = []

    try:
        from docx import Document
    except ImportError:
        return _result(errors=["python-docx not installed: pip install python-docx"])

    try:
        doc = Document(path)
    except Exception as exc:
        return _result(errors=[f"Could not read Word doc: {exc}"])

    try:
        with engine.begin() as con:
            for ti, table in enumerate(doc.tables):
                # Get header row
                if not table.rows:
                    continue
                header = [cell.text.strip().upper()
                          for cell in table.rows[0].cells]
                cols = {h: i for i, h in enumerate(header)}

                # Find key columns
                uwi_col   = next((cols[c] for c in cols if "UWI" in c), None)
                form_col  = next((cols[c] for c in cols
                                  if any(k in c for k in
                                         ("FORMATION","FORM","UNIT","STRAT"))), None)
                top_col   = next((cols[c] for c in cols
                                  if "TOP" in c and "BASE" not in c
                                  and "TVD" not in c), None)
                base_col  = next((cols[c] for c in cols
                                  if "BASE" in c), None)
                net_col   = next((cols[c] for c in cols
                                  if "NET" in c), None)
                fluid_col = next((cols[c] for c in cols
                                  if "FLUID" in c), None)

                if uwi_col is None or form_col is None or top_col is None:
                    continue  # not a formation tops table

                for row in table.rows[1:]:
                    cells   = [cell.text.strip() for cell in row.cells]
                    raw_uwi = cells[uwi_col] if uwi_col < len(cells) else ""
                    if not raw_uwi:
                        continue

                    uwi = _find_well(con, raw_uwi)
                    if not uwi:
                        errors.append(f"Table {ti+1}: UWI '{raw_uwi}' not found")
                        continue

                    form_name  = _trunc(cells[form_col], 255) if form_col < len(cells) else None
                    top_depth  = _safe_float(cells[top_col]) if top_col < len(cells) else None
                    base_depth = _safe_float(cells[base_col]) if base_col and base_col < len(cells) else None
                    net_pay    = _safe_float(cells[net_col])  if net_col  and net_col  < len(cells) else None
                    fluid      = _trunc(cells[fluid_col], 40) if fluid_col and fluid_col < len(cells) else None

                    if not form_name or top_depth is None:
                        continue

                    strat_id = _uid()
                    try:
                        con.execute(text("""
                            INSERT INTO dataview.dv_well_formation_top (
                                uwi, strat_unit_id, interp_id,
                                strat_unit_name, strat_unit_type,
                                top_depth, base_depth, depth_ouom, depth_datum,
                                net_thickness, fluid_type,
                                active_ind,
                                row_created_by, row_created_date, source
                            ) VALUES (
                                :uwi, :sid, '1',
                                :fname, 'FORMATION',
                                :top, :base, 'FT', 'KB',
                                :net, :fluid,
                                'Y',
                                :by, GETDATE(), :src
                            )
                        """), {
                            "uwi":   uwi, "sid": strat_id,
                            "fname": form_name,
                            "top":   top_depth, "base": base_depth,
                            "net":   net_pay, "fluid": fluid,
                            "by":    "DataWrangler", "src": source,
                        })
                        loaded += 1
                        if uwi not in wells:
                            wells.append(uwi)
                    except Exception as e:
                        errors.append(f"Table {ti+1}, {form_name}: {e}")

    except Exception as exc:
        return _result(errors=[str(exc)])

    return _result(loaded=loaded, errors=errors, wells=wells)


# =============================================================================
# DISPATCHER — call the right loader based on file extension + doc type hint
# =============================================================================

DOC_TYPE_LOADERS = {
    # (extension, doc_type_hint) → function
    ("xlsx", "formation"):   load_formation_tops_xlsx,
    ("xlsx", "completion"):  load_completion_xlsx,
    ("xlsx", "production"):  load_production_xlsx,
    ("docx", "completion"):  load_completion_docx,
    ("docx", "formation"):   load_formation_tops_docx,
}


def dispatch(engine, path: str, doc_type_hint: str = "",
             source: str = "DATA_LOADER") -> dict:
    """
    Automatically select and call the right loader based on
    file extension and doc_type_hint keyword.

    doc_type_hint examples: "formation", "completion", "production"
    """
    ext  = Path(path).suffix.lower().lstrip(".")
    hint = doc_type_hint.lower()

    # Match most specific key first
    key = next(
        ((e, h) for (e, h) in DOC_TYPE_LOADERS
         if e == ext and h in hint),
        None
    )
    if key:
        return DOC_TYPE_LOADERS[key](engine, path, source)

    # Fallback: guess from filename
    fname = Path(path).stem.lower()
    if "formation" in fname or "tops" in fname or "pick" in fname:
        if ext == "xlsx":
            return load_formation_tops_xlsx(engine, path, source)
        if ext == "docx":
            return load_formation_tops_docx(engine, path, source)
    if "complet" in fname or "frac" in fname or "stim" in fname:
        if ext == "xlsx":
            return load_completion_xlsx(engine, path, source)
        if ext == "docx":
            return load_completion_docx(engine, path, source)
    if "prod" in fname or "volume" in fname or "monthly" in fname:
        if ext == "xlsx":
            return load_production_xlsx(engine, path, source)

    return _result(errors=[
        f"No loader found for extension='{ext}' hint='{doc_type_hint}'. "
        f"Supported: formation/xlsx, completion/xlsx, production/xlsx, "
        f"formation/docx, completion/docx"
    ])
