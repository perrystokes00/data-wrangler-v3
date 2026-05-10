"""
modules/file_summarizer.py
==========================
Universal file summarizer for all Data Wrangler supported formats.
Returns a standard summary dict for any file type.

Usage:
    from modules.file_summarizer import summarize
    info = summarize("/path/to/well.las")
    print(info["description"])
"""
from __future__ import annotations
import re, uuid
from pathlib import Path
from typing import Optional

# ── Standard summary structure ────────────────────────────────────────────────
def _base(file_path: str, fmt: str) -> dict:
    fp = Path(file_path)
    return {
        "file_path":   file_path,
        "file_name":   fp.name,
        "format":      fmt,
        "size_kb":     round(fp.stat().st_size / 1024, 1) if fp.exists() else 0,
        "well_name":   None,
        "uwi":         None,
        "description": "",
        "key_fields":  {},
        "warnings":    [],
        "ppdm_hints":  [],
        "error":       None,
    }


# ══════════════════════════════════════════════════════════════════════════════
# LAS
# ══════════════════════════════════════════════════════════════════════════════
def _summarize_las(file_path: str) -> dict:
    s = _base(file_path, "LAS")
    try:
        import lasio
        las = lasio.read(file_path, ignore_header_errors=True)

        well    = las.well
        curves  = las.curves
        depth   = las.index

        s["well_name"] = (well.get("WELL") or well.get("WN") or
                          well.get("WELLNAME","")).value if hasattr(
                          well.get("WELL",""), "value") else str(
                          well.get("WELL",""))
        s["uwi"]       = str(well.get("UWI","") or well.get("API","") or "")

        n_curves  = len(curves)
        curve_nms = [c.mnemonic for c in curves
                     if c.mnemonic.upper() not in ("DEPT","DEPTH","MD")][:8]
        d_start   = round(float(depth.min()), 1) if len(depth) else 0
        d_stop    = round(float(depth.max()), 1) if len(depth) else 0
        d_step    = round(float(las.well.get("STEP","").value), 3) if hasattr(
                    las.well.get("STEP",""), "value") else 0
        null_val  = str(las.well.get("NULL","").value) if hasattr(
                    las.well.get("NULL",""), "value") else "-9999.25"
        company   = str(las.well.get("COMP","").value) if hasattr(
                    las.well.get("COMP",""), "value") else ""
        field     = str(las.well.get("FLD","").value) if hasattr(
                    las.well.get("FLD",""), "value") else ""

        s["description"] = (
            f"LAS {las.version.get('VERS','2.0').value if hasattr(las.version.get('VERS',''), 'value') else '2.0'}"
            f" · {n_curves} curves · {d_start:,.0f}–{d_stop:,.0f} ft"
            f" · {d_step} ft step"
            f" · Curves: {', '.join(curve_nms)}"
        )
        s["key_fields"] = {
            "curves":      n_curves,
            "curve_names": curve_nms,
            "depth_start": d_start,
            "depth_stop":  d_stop,
            "depth_step":  d_step,
            "null_value":  null_val,
            "company":     company,
            "field":       field,
            "samples":     len(depth),
        }
        s["ppdm_hints"] = ["dbo.WELL_LOG_SAMPLE", "dbo.WELL"]

        # Warnings
        if d_step == 0:
            s["warnings"].append("Step size is zero or missing")
        if not s["uwi"].strip():
            s["warnings"].append("No UWI/API found in header")
        if n_curves < 2:
            s["warnings"].append("Only 1 curve — file may be incomplete")

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# DLIS
# ══════════════════════════════════════════════════════════════════════════════
def _summarize_dlis(file_path: str) -> dict:
    s = _base(file_path, "DLIS")
    try:
        import dlisio
        with dlisio.dlis(file_path) as (f, *tail):
            lfs = [f] + list(tail)
            origins = f.origins
            if origins:
                o = origins[0]
                s["well_name"] = str(getattr(o, "well_name", "") or "")
                s["uwi"]       = str(getattr(o, "api_well", "") or
                                     getattr(o, "uwi", "") or "")
                company        = str(getattr(o, "company", "") or "")
                field          = str(getattr(o, "field_name", "") or "")
            else:
                company = field = ""

            total_ch = sum(len(lf.channels) for lf in lfs)
            total_fr = sum(len(lf.frames)   for lf in lfs)

            s["description"] = (
                f"DLIS · {len(lfs)} logical file(s) · "
                f"{total_ch} channels · {total_fr} frame(s)"
            )
            s["key_fields"] = {
                "logical_files": len(lfs),
                "channels":      total_ch,
                "frames":        total_fr,
                "company":       company,
                "field":         field,
            }
            s["ppdm_hints"] = ["dbo.WELL_LOG_SAMPLE", "dbo.WELL"]

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# SEG-Y
# ══════════════════════════════════════════════════════════════════════════════
def _summarize_segy(file_path: str) -> dict:
    s = _base(file_path, "SEG-Y")
    try:
        import segyio
        with segyio.open(file_path, ignore_geometry=True) as f:
            n_traces  = f.tracecount
            si        = segyio.dt(f) / 1000.0   # sample interval ms
            n_samples = f.samples.size
            rec_len   = round(si * n_samples, 0)
            fmt_code  = f.bin[segyio.BinField.DataSampleFormat]
            fmt_names = {1:"IBM float",2:"Int32",3:"Int16",
                         4:"Fixed pt",5:"IEEE float",8:"Int8"}
            fmt_name  = fmt_names.get(fmt_code, str(fmt_code))

            # Try to get inline/crossline range
            try:
                il_min = int(min(f.attributes(segyio.TraceField.INLINE_3D)[:]))
                il_max = int(max(f.attributes(segyio.TraceField.INLINE_3D)[:]))
                xl_min = int(min(f.attributes(segyio.TraceField.CROSSLINE_3D)[:]))
                xl_max = int(max(f.attributes(segyio.TraceField.CROSSLINE_3D)[:]))
                geom = f"IL {il_min}-{il_max} · XL {xl_min}-{xl_max}"
            except Exception:
                geom = "geometry not available"

            s["description"] = (
                f"SEG-Y · {n_traces:,} traces · {si}ms sample interval"
                f" · {rec_len:.0f}ms record · {fmt_name} · {geom}"
            )
            s["key_fields"] = {
                "traces":          n_traces,
                "sample_interval": si,
                "record_length_ms":rec_len,
                "samples":         n_samples,
                "format":          fmt_name,
                "format_code":     fmt_code,
            }
            s["ppdm_hints"] = ["dbo.SEIS_SET", "dbo.SEIS_TRACE"]

            if n_traces == 0:
                s["warnings"].append("Zero traces — file may be corrupt")

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════════════════════
def _summarize_pdf(file_path: str) -> dict:
    s = _base(file_path, "PDF")
    try:
        import pdfplumber
        from modules.pdf_survey_catalog import classify_pdf

        cl = classify_pdf(file_path)
        s["well_name"] = cl.get("well_name")
        s["uwi"]       = cl.get("uwi")

        with pdfplumber.open(file_path) as pdf:
            pages    = len(pdf.pages)
            has_tbl  = any(pdf.pages[i].extract_tables()
                           for i in range(min(3, pages)))

        s["description"] = (
            f"PDF · {pages} page(s) · {cl['report_type'].replace('_',' ').title()}"
            f" · {cl.get('station_count',0)} stations"
            f" · {int(cl.get('confidence',0)*100)}% confidence"
        )
        s["key_fields"] = {
            "pages":        pages,
            "report_type":  cl["report_type"],
            "station_count":cl.get("station_count", 0),
            "confidence":   cl.get("confidence", 0),
            "has_tables":   has_tbl,
            "operator":     cl.get("operator"),
            "survey_type":  cl.get("survey_type"),
        }
        if cl["report_type"] == "DIRECTIONAL_SURVEY":
            s["ppdm_hints"] = ["dbo.WELL_DIR_SURVEY",
                               "dbo.WELL_DIR_SRVY_STATION"]
        elif cl["report_type"] == "MUD_LOG":
            s["ppdm_hints"] = ["dbo.WELL_LOG_SAMPLE"]

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# Shapefile
# ══════════════════════════════════════════════════════════════════════════════
def _summarize_shp(file_path: str) -> dict:
    s = _base(file_path, "Shapefile")
    try:
        import geopandas as gpd
        from modules.shapefile_catalog import classify_shapefile

        cl = classify_shapefile(file_path)
        gdf = gpd.read_file(file_path, rows=1)

        n       = cl.get("feature_count", 0)
        geom    = cl.get("geometry_type", "?")
        crs     = cl.get("crs_epsg", "?")
        cols    = cl.get("attributes", [])
        ft      = cl.get("feature_type", "?")
        bounds  = cl.get("bounds") or {}

        s["description"] = (
            f"Shapefile · {n:,} {geom} features · {ft.replace('_',' ')}"
            f" · CRS EPSG:{crs} · {len(cols)} attributes"
        )
        s["key_fields"] = {
            "feature_count": n,
            "geometry_type": geom,
            "feature_type":  ft,
            "crs_epsg":      crs,
            "attributes":    cols[:10],
            "bounds":        bounds,
            "confidence":    cl.get("confidence", 0),
        }
        s["ppdm_hints"] = [cl.get("ppdm_target")] if cl.get("ppdm_target") else []

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# Excel
# ══════════════════════════════════════════════════════════════════════════════

# Column pattern → PPDM table hints
EXCEL_TABLE_TYPES = {
    "PRODUCTION": {
        "keywords": ["oil","gas","water","bbls","mcf","boe",
                     "production","prod","volume","gross","net"],
        "ppdm":     "dbo.WELL_VERSION",
        "required": ["date","uwi"],
    },
    "COMPLETION": {
        "keywords": ["frac","stage","proppant","fluid","perf",
                     "cluster","completion","sand","water_vol"],
        "ppdm":     "dbo.WELL_COMPLETION",
        "required": ["uwi"],
    },
    "FORMATION_TOPS": {
        "keywords": ["formation","top","base","net_pay","pay",
                     "tops","horizon","marker","depth"],
        "ppdm":     "dbo.WELL_FORMATION",
        "required": ["uwi","formation"],
    },
    "WELL_HEADER": {
        "keywords": ["uwi","api","well_name","operator","latitude",
                     "longitude","field","county","kb","td"],
        "ppdm":     "dbo.WELL",
        "required": ["uwi"],
    },
    "CORE_ANALYSIS": {
        "keywords": ["porosity","permeability","perm","poro","sw",
                     "saturation","core","plug","grain"],
        "ppdm":     "dbo.WELL_CORE_ANALYSIS",
        "required": ["depth"],
    },
    "PRESSURE": {
        "keywords": ["pressure","psi","bhp","whp","temperature",
                     "gradient","datum","shut_in"],
        "ppdm":     "dbo.WELL_PRESSURE_SURVEY",
        "required": ["depth","pressure"],
    },
    "SURVEY": {
        "keywords": ["md","inc","azi","tvd","inclination","azimuth",
                     "measured_depth","dogleg","dls"],
        "ppdm":     "dbo.WELL_DIR_SRVY_STATION",
        "required": ["md","inc"],
    },
    "RESERVE": {
        "keywords": ["reserve","proved","probable","possible",
                     "1p","2p","3p","pdp","pud","mstb","mmcf"],
        "ppdm":     "dbo.WELL_VERSION",
        "required": ["uwi"],
    },
}


def _classify_excel_sheet(headers: list[str]) -> tuple[str, float]:
    """Classify a sheet by its column headers."""
    hdrs_lower = [h.lower().replace(' ','_').replace('/','_')
                  for h in headers if h]
    best_type  = "UNKNOWN"
    best_score = 0.0

    for table_type, cfg in EXCEL_TABLE_TYPES.items():
        score = sum(
            1 for kw in cfg["keywords"]
            if any(kw in h for h in hdrs_lower)
        )
        req_score = sum(
            1 for req in cfg["required"]
            if any(req in h for h in hdrs_lower)
        )
        # Weight required columns more heavily
        total = score + req_score * 2
        norm  = total / (len(cfg["keywords"]) + len(cfg["required"]) * 2)
        if norm > best_score:
            best_score = norm
            best_type  = table_type

    return best_type, round(best_score, 2)


def _summarize_excel(file_path: str) -> dict:
    s = _base(file_path, "Excel")
    try:
        import openpyxl
        import pandas as pd

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheets = wb.sheetnames

        sheet_summaries = []
        all_ppdm        = []
        total_rows      = 0
        uwi_found       = None

        for sheet_name in sheets:
            try:
                # Read just headers and first few rows
                df = pd.read_excel(
                    file_path, sheet_name=sheet_name,
                    nrows=5, engine='openpyxl'
                )
                if df.empty or len(df.columns) < 2:
                    continue

                headers      = [str(c) for c in df.columns]
                table_type, conf = _classify_excel_sheet(headers)

                # Get full row count
                ws        = wb[sheet_name]
                n_rows    = ws.max_row - 1  # subtract header
                total_rows += max(0, n_rows)

                # Try to find UWI in data
                if not uwi_found:
                    for col in headers:
                        if any(x in col.lower() for x in
                               ['uwi','api','well_id']):
                            try:
                                df2 = pd.read_excel(
                                    file_path, sheet_name=sheet_name,
                                    nrows=2, engine='openpyxl'
                                )
                                if not df2[col].empty:
                                    uwi_found = str(df2[col].iloc[0])
                            except Exception:
                                pass

                # Date range for production/time series
                date_range = ""
                for col in headers:
                    if any(x in col.lower() for x in ['date','month','year']):
                        try:
                            df3 = pd.read_excel(
                                file_path, sheet_name=sheet_name,
                                usecols=[col], engine='openpyxl'
                            )
                            dates = pd.to_datetime(df3[col],
                                    errors='coerce').dropna()
                            if len(dates) > 0:
                                date_range = (
                                    f"{dates.min().strftime('%Y-%m')} – "
                                    f"{dates.max().strftime('%Y-%m')}"
                                )
                        except Exception:
                            pass
                        break

                ppdm = EXCEL_TABLE_TYPES.get(
                    table_type, {}).get("ppdm","")
                if ppdm:
                    all_ppdm.append(ppdm)

                sheet_summaries.append({
                    "sheet":      sheet_name,
                    "table_type": table_type,
                    "confidence": conf,
                    "rows":       n_rows,
                    "columns":    len(headers),
                    "headers":    headers[:8],
                    "ppdm":       ppdm,
                    "date_range": date_range,
                })

            except Exception:
                pass

        wb.close()

        s["uwi"] = uwi_found
        s["description"] = (
            f"Excel · {len(sheets)} sheet(s) · {total_rows:,} total rows · "
            + " | ".join(
                f"{ss['sheet']}: {ss['table_type']} ({ss['rows']:,} rows)"
                for ss in sheet_summaries[:4]
            )
        )
        s["key_fields"] = {
            "sheets":       sheets,
            "total_rows":   total_rows,
            "sheet_detail": sheet_summaries,
        }
        s["ppdm_hints"] = list(dict.fromkeys(all_ppdm))  # dedup, keep order

        if total_rows == 0:
            s["warnings"].append("No data rows found — file may be empty")
        if not s["ppdm_hints"]:
            s["warnings"].append("Could not classify sheet content")

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# Word
# ══════════════════════════════════════════════════════════════════════════════

WORD_DOC_TYPES = {
    "COMPLETION_REPORT":  ["completion","perforation","frac","stimulation",
                           "proppant","wellbore"],
    "GEOLOGICAL_REPORT":  ["formation","geology","geological","lithology",
                           "stratigraph","pay zone","net pay"],
    "DST_REPORT":         ["drill stem test","dst","pressure buildup",
                           "shut-in","flow rate","bhp","isip"],
    "WELL_PROPOSAL":      ["afe","authorization","proposed well",
                           "prognosis","objective"],
    "REGULATORY":         ["permit","regulatory","state","commission",
                           "railroad","compliance","notif"],
    "FORMATION_TOPS":     ["formation tops","top of","base of",
                           "picked at","marker"],
    "HSE_REPORT":         ["safety","incident","hse","near miss",
                           "hazard","injury"],
}


def _summarize_docx(file_path: str) -> dict:
    s = _base(file_path, "Word")
    try:
        import docx

        doc     = docx.Document(file_path)
        core    = doc.core_properties

        # Document properties
        title   = core.title   or ""
        author  = core.author  or ""
        created = str(core.created)[:10] if core.created else ""
        modified= str(core.modified)[:10] if core.modified else ""

        # Extract all text
        full_text = "\n".join(p.text for p in doc.paragraphs)
        text_up   = full_text.upper()

        # Extract headings
        headings = [p.text.strip() for p in doc.paragraphs
                    if p.style.name.startswith("Heading")
                    and p.text.strip()][:10]

        # Count tables and extract headers
        tables_info = []
        for i, tbl in enumerate(doc.tables[:10]):
            if not tbl.rows:
                continue
            hdr_row  = [c.text.strip() for c in tbl.rows[0].cells]
            n_rows   = len(tbl.rows) - 1
            tbl_type, conf = _classify_excel_sheet(hdr_row)
            tables_info.append({
                "table_idx":  i,
                "headers":    hdr_row[:8],
                "rows":       n_rows,
                "table_type": tbl_type,
                "confidence": conf,
                "ppdm":       EXCEL_TABLE_TYPES.get(tbl_type,{}).get("ppdm",""),
            })

        # Detect document type
        doc_type = "UNKNOWN"
        best_score = 0
        for dtype, keywords in WORD_DOC_TYPES.items():
            score = sum(1 for kw in keywords if kw.upper() in text_up)
            if score > best_score:
                best_score = score
                doc_type   = dtype

        # Extract UWI/well name
        uwi_match = re.search(
            r'(?:UWI|API)[:\s]+([0-9\-]{10,20})', full_text, re.IGNORECASE)
        well_match = re.search(
            r'(?:WELL\s+NAME|WELL)[:\s]+([A-Za-z0-9 #\-]+)',
            full_text, re.IGNORECASE)

        s["uwi"]       = uwi_match.group(1).strip() if uwi_match else None
        s["well_name"] = well_match.group(1).strip()[:50] if well_match else None

        ppdm_hints = list(dict.fromkeys(
            t["ppdm"] for t in tables_info if t.get("ppdm")
        ))

        s["description"] = (
            f"Word · {doc_type.replace('_',' ').title()}"
            f" · {len(doc.paragraphs)} paragraphs"
            f" · {len(doc.tables)} table(s)"
            f" · {len(full_text):,} characters"
        )
        s["key_fields"] = {
            "doc_type":     doc_type,
            "title":        title,
            "author":       author,
            "created":      created,
            "modified":     modified,
            "headings":     headings,
            "paragraphs":   len(doc.paragraphs),
            "tables":       len(doc.tables),
            "tables_detail":tables_info,
            "word_count":   len(full_text.split()),
        }
        s["ppdm_hints"] = ppdm_hints

        if not doc.paragraphs:
            s["warnings"].append("Document appears empty")
        if best_score == 0:
            s["warnings"].append("Could not classify document type")

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# CSV
# ══════════════════════════════════════════════════════════════════════════════
def _summarize_csv(file_path: str) -> dict:
    s = _base(file_path, "CSV")
    try:
        import pandas as pd

        # Read just headers + sample
        df   = pd.read_csv(file_path, nrows=5, low_memory=False)
        hdrs = [str(c) for c in df.columns]

        # Full row count without reading all data
        with open(file_path, 'r', errors='ignore') as f:
            n_rows = sum(1 for _ in f) - 1  # subtract header

        table_type, conf = _classify_excel_sheet(hdrs)
        ppdm = EXCEL_TABLE_TYPES.get(table_type, {}).get("ppdm","")

        # Find UWI
        for col in hdrs:
            if any(x in col.lower() for x in ['uwi','api','well_id']):
                if not df[col].empty:
                    s["uwi"] = str(df[col].iloc[0])
                break

        s["description"] = (
            f"CSV · {n_rows:,} rows · {len(hdrs)} columns · "
            f"{table_type.replace('_',' ').title()} ({int(conf*100)}%)"
        )
        s["key_fields"] = {
            "rows":       n_rows,
            "columns":    len(hdrs),
            "headers":    hdrs[:12],
            "table_type": table_type,
            "confidence": conf,
            "sample":     df.head(3).to_dict("records"),
        }
        s["ppdm_hints"] = [ppdm] if ppdm else []

        if n_rows == 0:
            s["warnings"].append("CSV has no data rows")

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# P190
# ══════════════════════════════════════════════════════════════════════════════
def _summarize_p190(file_path: str) -> dict:
    s = _base(file_path, "P190")
    try:
        lines = []
        with open(file_path, 'r', errors='ignore') as f:
            lines = f.readlines()[:200]

        header_lines = [l for l in lines if l.startswith('H')]
        data_lines   = [l for l in lines
                        if l[0:1].upper() in ('S','T','C')]

        survey_name = ""
        vessel      = ""
        for hl in header_lines:
            if 'SURVEY' in hl.upper():
                survey_name = hl[2:].strip()[:60]
            if 'VESSEL' in hl.upper() or 'SHIP' in hl.upper():
                vessel = hl[2:].strip()[:40]

        # Full count
        with open(file_path, 'r', errors='ignore') as f:
            all_lines  = f.readlines()
        total_data = sum(1 for l in all_lines
                         if l[0:1].upper() in ('S','T','C'))

        s["description"] = (
            f"P190 · {len(header_lines)} header records"
            f" · {total_data:,} data records"
            f" · Survey: {survey_name or 'unknown'}"
        )
        s["key_fields"] = {
            "header_records": len(header_lines),
            "data_records":   total_data,
            "survey_name":    survey_name,
            "vessel":         vessel,
        }
        s["ppdm_hints"] = ["dbo.SEIS_LINE", "dbo.SEIS_SET"]

    except Exception as e:
        s["error"] = str(e)
    return s


# ══════════════════════════════════════════════════════════════════════════════
# Main dispatcher
# ══════════════════════════════════════════════════════════════════════════════
def summarize(file_path: str) -> dict:
    """
    Universal summarizer — dispatches to the correct format handler.
    Always returns a dict with at minimum:
        file_path, file_name, format, size_kb,
        well_name, uwi, description, key_fields,
        warnings, ppdm_hints, error
    """
    ext = Path(file_path).suffix.lower()

    dispatch = {
        ".las":    _summarize_las,
        ".dlis":   _summarize_dlis,
        ".dlf":    _summarize_dlis,
        ".segy":   _summarize_segy,
        ".sgy":    _summarize_segy,
        ".seg":    _summarize_segy,
        ".pdf":    _summarize_pdf,
        ".shp":    _summarize_shp,
        ".geojson":_summarize_shp,
        ".gpkg":   _summarize_shp,
        ".xlsx":   _summarize_excel,
        ".xls":    _summarize_excel,
        ".docx":   _summarize_docx,
        ".doc":    _summarize_docx,
        ".csv":    _summarize_csv,
        ".txt":    _summarize_csv,
        ".tsv":    _summarize_csv,
        ".p190":   _summarize_p190,
        ".p90":    _summarize_p190,
        ".p1":     _summarize_p190,
    }

    handler = dispatch.get(ext)
    if handler:
        return handler(file_path)

    # Unknown format
    s = _base(file_path, ext.lstrip(".").upper() or "UNKNOWN")
    s["description"] = f"Unsupported format: {ext}"
    s["warnings"].append(f"No summarizer available for {ext} files")
    return s
